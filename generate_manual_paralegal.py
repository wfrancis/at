#!/usr/bin/env python3
"""Generate a paralegal-friendly Legal Department Operations Manual for PIR Industrial LLC.

Mirrors the structure of generate_manual.py (same 13 sections + Attachments A/B/C),
but rewrites prose in plain English so paralegals can digest it quickly.
Short sentences, active voice, fewer abstractions, same content coverage.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FONT = "Times New Roman"
FONT_SIZE = Pt(12)

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# -- Default style --
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = FONT_SIZE
style.paragraph_format.space_after = Pt(6)

# -- Header/Footer --
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = hp.add_run("PIR Industrial LLC \u2013 Legal Department Operations Manual (Paralegal Edition)")
run.font.name = FONT
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("Page ")
run.font.name = FONT
run.font.size = Pt(9)
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
run2 = fp.add_run()
run2._r.append(fldChar1)
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = ' PAGE '
run3 = fp.add_run()
run3._r.append(instrText)
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run4 = fp.add_run()
run4._r.append(fldChar2)


def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = FONT
    run.font.size = FONT_SIZE
    return p


def add_sub_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.name = FONT
    run.font.size = FONT_SIZE
    return p


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if "**" in text:
        parts = text.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            run.font.name = FONT
            run.font.size = FONT_SIZE
            if i % 2 == 1:
                run.bold = True
            if italic:
                run.italic = True
    else:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = FONT_SIZE
        if bold:
            run.bold = True
        if italic:
            run.italic = True
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        run = p.add_run(item)
        run.font.name = FONT
        run.font.size = FONT_SIZE


def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.clear()
        run = p.add_run(item)
        run.font.name = FONT
        run.font.size = FONT_SIZE


def add_table(headers, rows, col_widths_inches=None):
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = FONT
            run.font.size = Pt(10)

    if col_widths_inches:
        for row in table.rows:
            for i, width in enumerate(col_widths_inches):
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table


def page_break():
    doc.add_page_break()


# ============================================================
# HOW TO USE THIS MANUAL (paralegal intro)
# ============================================================
add_section_header("How to Use This Manual")
add_para(
    "This is the paralegal edition of the PIR Industrial LLC Legal Department Operations Manual. "
    "It covers the same 13 sections and 3 attachments as the full manual, but in plainer language."
)
add_para("Quick terminology:")
add_bullets([
    "\"Company\" means PIR Industrial LLC.",
    "\"Department\" or \"Legal\" means the Legal Department.",
    "\"Manual\" means this document.",
    "\"Outside counsel\" means law firms the Company hires.",
    "\"GC\" means General Counsel. \"CFO\" means Chief Financial Officer. \"CEO\" means Chief Executive Officer.",
    "\"REIT\" means Real Estate Investment Trust (a tax status the Company has elected).",
    "\"SharePoint\" is where we collaborate on drafts. \"ShareFile\" is where we store signed documents.",
])
add_para(
    "If something in this edition is unclear or conflicts with the full manual, the full manual controls. "
    "Ask the GC or your supervising attorney when in doubt."
)

# ============================================================
# SECTION 1: PURPOSE AND SCOPE
# ============================================================
page_break()
add_section_header("1. Purpose and Scope")
add_para(
    "This Manual tells Legal staff how to run legal matters at the Company. "
    "It covers what Legal does, how Legal works with other departments, and the steps to follow for common tasks."
)
add_para(
    "The goal is simple: give everyone in Legal a consistent playbook so business decisions get reliable legal support."
)
add_para(
    "The Company will update this Manual over time as tools, systems, and priorities change."
)

# ============================================================
# SECTION 2: OPERATING PRINCIPLES
# ============================================================
page_break()
add_section_header("2. Operating Principles")
add_para(
    "These principles guide how Legal works day to day. "
    "They cover how we balance speed with care, protect information, and work with other teams."
)

add_sub_header("Business Alignment")
add_para(
    "Legal supports the business. We give practical advice that helps deals get done while protecting the Company."
)
add_para(
    "Know the business: understand the Company's priorities and portfolio before giving advice. Balance risk against getting things done."
)

add_sub_header("Speed with Discipline")
add_para("Leasing moves fast. Legal has to move fast too, without cutting corners.")
add_para(
    "Respond quickly, but still check accuracy, track versions, and coordinate internally. Good workflows let us be fast and careful."
)

add_sub_header("Consistency and Standardization")
add_para(
    "Follow the same procedures every time. Consistency keeps work reliable and makes it easy for others to know what to expect."
)

add_sub_header("Transparency and Risk Visibility")
add_para("Leadership needs to see legal risks. Legal uses trackers and reports to make risks visible.")
add_para(
    "Spot risks early, write them down, and tell the right people in time. Clear reporting helps leaders decide and act."
)

add_sub_header("Centralized Document Management")
add_para(
    "Store every legal document in Company systems: SharePoint for drafts in progress, ShareFile for signed documents."
)
add_para(
    "Do not save files to your laptop or personal cloud drives. Keep folders organized and follow the naming rules."
)

add_sub_header("Effective Collaboration")
add_para("Legal works closely with Asset Management, Finance, Property Management, and executive leadership.")
add_para(
    "Communicate clearly, agree on timelines, and know who owns what. Legal advises; the business team owns the decision."
)

add_sub_header("Prudent Use of Outside Counsel")
add_para("Outside counsel fill gaps: specialized expertise, jurisdictions we don't cover, or extra capacity when we need it.")
add_para(
    "Legal picks the firm, sets the scope, watches the work, and manages the bill. "
    "Outside counsel should act like part of our team and follow our billing guidelines."
)

add_sub_header("Continuous Improvement")
add_para("Keep looking for better ways to work: new tools, cleaner workflows, better templates.")
add_para("Evaluate technology that helps with documents, reporting, and data. Keep training materials up to date.")
add_para("Improvement keeps Legal effective as the Company grows.")

add_sub_header("Accountability")
add_para("Every member of Legal is expected to use good judgment, act with integrity, and own their work.")
add_para("Keep confidences, follow the law, follow Company policies, and be honest about what you know and don't know.")

# ============================================================
# SECTION 3: ROLES AND RESPONSIBILITIES
# ============================================================
page_break()
add_section_header("3. Roles and Responsibilities")
add_para(
    "Legal provides legal advice, drafts documents, and oversees risk. "
    "On some topics Legal is in charge; on others Legal advises and another department decides."
)
add_para("The lists below show what Legal owns, what Legal advises on, and what Legal stays out of.")

add_sub_header("What Legal Owns")

add_para("**Lease Documents and Drafting**")
add_bullets([
    "Draft and review leases and amendments.",
    "Keep standard lease forms current.",
    "Make sure documents match the approved business deal and protect the Company.",
    "Manage internal review and keep track of versions.",
])

add_para("**Litigation Oversight**")
add_bullets([
    "Run the Company's lawsuits.",
    "Pick and supervise outside counsel.",
    "Set litigation strategy and settlement position with executive leadership.",
    "Watch the budget and control spend.",
])

add_para("**Outside Counsel Management**")
add_bullets([
    "Maintain a panel of approved firms.",
    "Send engagement letters and set the scope.",
    "Review firm performance, staffing, and bills.",
    "Enforce the Company's billing guidelines and reporting requirements.",
])

add_para("**Legal Risk Tracking**")
add_bullets([
    "Maintain the Legal Risk Tracker.",
    "Identify risks across the portfolio.",
    "Report material matters to executive leadership.",
    "Recommend how to reduce risk.",
])

add_para("**Corporate Governance and Entities**")
add_bullets([
    "Keep entity records current.",
    "Run governance processes for the Company and its subsidiaries.",
    "Handle entity formation, dissolution, and restructurings.",
    "Maintain org charts and ownership records.",
])

add_para("**Policy and Compliance**")
add_bullets([
    "Write and update internal legal policies.",
    "Watch for new laws and regulations that affect the Company.",
    "Advise other teams on compliance.",
])

add_sub_header("Where Legal Advises")

add_para("**Lease Negotiation (with Asset Management)**")
add_bullets([
    "Flag legal implications of proposed terms.",
    "Suggest ways to reduce risk.",
    "Review tenant revisions.",
])

add_para("**Tenant Defaults and Disputes (with Executive)**")
add_bullets([
    "Default notice steps.",
    "Options for enforcement.",
    "Settlement considerations.",
    "Litigation strategy if it escalates.",
])

add_para("**Insurance Claims (with Asset Management)**")
add_bullets([
    "Complicated coverage questions.",
    "Disputes with insurers.",
    "Lawsuits tied to claims.",
])

add_para("**Transactional Work (with Executive, Finance, and Acquisitions)**")
add_bullets([
    "Buying and selling properties.",
    "Financing.",
    "Joint ventures.",
    "Material contracts.",
])

add_sub_header("What Legal Does Not Own")
add_bullets([
    "Relationships with tenants, vendors, contractors, and property managers.",
    "Rent collection and accounts receivable.",
    "Property operations and maintenance.",
    "Leasing strategy and market positioning.",
    "Asset Management decisions about tenant concessions or pricing.",
])

# ============================================================
# SECTION 4: INTAKE & PRIORITIZATION
# ============================================================
page_break()
add_section_header("4. Intake and Prioritization")
add_para(
    "Legal gets requests from Asset Management, Property Management, Finance, and executive leadership. "
    "Every request goes through a standard intake so nothing falls through the cracks."
)
add_para("Why we do it this way:")
add_numbered([
    "Every request is logged and tracked.",
    "Priorities are set by clear rules.",
    "Leadership can see how busy Legal is.",
    "Urgent matters (deadlines, revenue, legal rights) get attention first.",
])

add_sub_header("How to Submit a Legal Request")
add_para("All leasing requests go through Salesforce (the Company's lease lifecycle platform).")
add_para("Salesforce automatically tells Legal when Asset Management asks for:")
add_bullets([
    "A new lease or amendment draft.",
    "Review of a tenant's draft or amendment.",
    "A tenant default notice.",
    "Review of any other property document.",
])
add_para("Using Salesforce time-stamps the request, assigns an owner, and tracks it to the finish.")
add_para("When submitting, Asset Management enters:")
add_bullets([
    "Property name and address.",
    "Tenant name.",
    "Document type.",
    "Whether this is a first draft or a tenant revision.",
    "Lease or renewal start date (if any).",
    "Total rent over the lease term.",
])
add_para("This info tells Legal how to prioritize and staff the request.")

add_sub_header("How We Prioritize")
add_para(
    "The head of Asset Management (and other authorized people) has the final say on lease priority. "
    "Over time, Salesforce should automate prioritization using the info above. Priority factors include:"
)
add_bullets([
    "Leases starting soon.",
    "Big rent or strategic tenants.",
    "Active tenant revisions (to keep the deal moving).",
    "Hard-to-lease or strategically important properties.",
    "Contractual, regulatory, or litigation deadlines.",
])

add_para("**High priority:**")
add_bullets([
    "New leases with a rent start date coming up.",
    "Tenant revisions that need a fast turnaround.",
    "Big revenue impact.",
    "Legal deadlines or risk of losing a right by not acting.",
    "Anything executive leadership escalates.",
])
add_para("Aim to respond in 1 to 2 business days, depending on complexity.")

add_para("**Medium priority:**")
add_bullets([
    "Amendments with no immediate start date.",
    "Ongoing negotiations with no hard deadline.",
    "Routine reviews (landlord lien waivers, NDAs, collateral access agreements, SNDAs, estoppels, etc.).",
])
add_para("These go through the normal queue.")

add_para("**Low priority:**")
add_bullets([
    "Housekeeping amendments.",
    "Internal document reviews.",
    "Information-only requests.",
    "Long-term projects with no time pressure.",
])
add_para("Handle these when you have capacity.")

add_sub_header("Legal Workflow Dashboard")
add_para(
    "Salesforce rolls up all legal requests into one dashboard. "
    "It shows pending work, priority, owner, draft status, and expected finish date."
)
add_para("The dashboard gives leadership a view of the pipeline and how work is flowing.")

add_sub_header("Reviewing Priorities")
add_para("Legal and Asset Management should check open matters regularly to make sure priorities still fit the situation.")
add_para(
    "Priorities usually follow the rules above, but sometimes business judgment calls for a change. "
    "Authorized people can override the standard priority when the Company's needs require it."
)

# ============================================================
# SECTION 5: LEASING ASSET MANAGEMENT PROCEDURES
# ============================================================
page_break()
add_section_header("5. Leasing Procedures")
add_para(
    "Leasing is a core part of the Company's business. "
    "Legal supports Asset Management on lease negotiations and documents and makes sure the Company is protected."
)
add_para("This section covers:")
add_bullets([
    "How leases get drafted and reviewed.",
    "How Legal and Asset Management work together.",
    "How to manage document versions.",
    "How drafts get circulated internally and to the tenant.",
])
add_para("The goal is accurate documents, smooth collaboration, and clear ownership.")

add_sub_header("Leasing Workflow: The Steps")
add_para("A typical lease moves through these steps:")
add_numbered([
    "Asset Management sets up a SharePoint workspace for the deal. Draft documents (lease analysis, LOIs, etc.) go there for the deal team to review.",
    "Asset Management submits the drafting or review request in Salesforce's Leasing Process system.",
    "Legal drafts or reviews the document. If it's a tenant's draft, Asset Management first adds its comments and tracked changes in Word. Legal then edits and adds comments using the Word \"@mention\" feature. Use in-document comments, not email.",
    "When internal review is done, Legal or Asset Management runs Word's \"Compare\" feature against the tenant's last version (if any), downloads a copy from SharePoint, and sends it to tenant counsel or the tenant. Asset Management updates Salesforce to show the deal is in negotiation.",
    "Drafts get revised as negotiations continue. Save every internal and external version as a separate Word file using the naming rule below.",
    "When the document is final, the Leasing Administrator in Asset Management handles signatures and saves the executed document in ShareFile (not SharePoint).",
])

add_sub_header("What Asset Management Does")
add_para("Asset Management owns the business negotiation and the tenant or broker relationship.")
add_para("Asset Management's job:")
add_bullets([
    "Negotiate economic terms.",
    "Work with the tenant and broker on business points.",
    "Pass drafts between Legal, tenant, and broker.",
    "Check that the final document matches the negotiated deal.",
])
add_para(
    "Sometimes Asset Management drafts a business term directly, but Legal still has final sign-off. "
    "If timing doesn't allow Legal to review before sending, Asset Management can transmit the document, "
    "but must say in writing that Legal still has the right to review and revise."
)

add_para("Legal's job:")
add_bullets([
    "Draft leases and amendments.",
    "Review tenant-proposed forms.",
    "Keep documents consistent with the Company's standard provisions.",
    "Explain the legal impact of tenant or Asset Management revisions.",
    "Track versions.",
    "Coordinate internal review before anything goes to the tenant.",
])

add_sub_header("SharePoint (Drafts in Progress)")
add_para("Every lease document under negotiation lives in that deal's SharePoint workspace.")
add_para("SharePoint gives us:")
add_bullets([
    "One place to store files.",
    "Version history for an audit trail.",
    "Shared editing with the deal team.",
    "Document integrity (no stray copies).",
])
add_para("Centralizing documents means everyone works from the same current file.")

add_sub_header("Document Naming")
add_para("Name every draft the same way so the team always knows what's what.")
add_para("Format:")
add_para("**[Document Type] \u2013 [Tenant Name] \u2013 [Property Address] \u2013 PLYM Draft**")
add_para(
    "Number each version in order. The tenant's first proposed form is version 1 for internal review. "
    "The highest-numbered version is the current draft for sending out or for signature."
)
add_para("Lease \u2013 ABC Logistics \u2013 100 Industrial Drive \u2013 PLYM Draft v2", italic=True)
add_para("Lease \u2013 ABC Logistics \u2013 100 Industrial Drive \u2013 PLYM Draft v3", italic=True)
add_para("Consistent names prevent mix-ups and make the latest version easy to spot.")

add_sub_header("Tracked Changes")
add_para(
    "Keep tracked changes on during negotiation so everyone can see what changed. "
    "Do not accept or strip changes before internal review is done. "
    "Before sending the draft to the tenant, resolve internal comments, delete them, then run a Word \"Compare\" "
    "against the tenant's last version so the recipient can see the diff clearly."
)
add_para("Tracked changes keep the negotiation transparent and efficient.")

add_sub_header("Improving Lease Forms Over Time")
add_para("Legal periodically reviews leasing activity to improve standard forms and negotiation habits.")
add_para("Review topics include:")
add_bullets([
    "Terms tenants push back on most.",
    "Recurring tenant concerns.",
    "New laws affecting lease provisions.",
    "Feedback from Asset Management.",
])
add_para("These reviews keep our documents current with the market, operations, and risk.")

# ============================================================
# SECTION 6: LITIGATION & OUTSIDE COUNSEL MANAGEMENT
# ============================================================
page_break()
add_section_header("6. Litigation and Outside Counsel")
add_para(
    "Legal runs all Company litigation and disputes. That means managing risk, picking and supervising outside counsel, "
    "controlling legal spend, and keeping strategy in line with business goals."
)
add_para("The procedures below:")
add_bullets([
    "Keep litigation centrally managed.",
    "Give leadership a clear picture of legal exposure.",
    "Keep outside counsel cost-effective.",
    "Support timely, informed decisions.",
])

add_sub_header("Logging New Matters")
add_para(
    "The moment we get a demand letter, complaint, or notice of claim, log it in the Legal Risk Tracker. "
    "Legal assigns a unique ID and names a responsible attorney or paralegal."
)
add_para("For each matter, record:")
add_bullets([
    "Who is involved.",
    "What the claim or dispute is about.",
    "Which property or entity it affects.",
    "Date received and key deadlines.",
    "Outside counsel (if any).",
    "Estimated financial exposure.",
    "Strategy and settlement position.",
])

add_sub_header("Hiring Outside Counsel")
add_para(
    "Legal keeps an approved panel of firms chosen for expertise, jurisdiction coverage, responsiveness, and cost. "
    "Only Legal hires outside counsel."
)
add_para("Before hiring, Legal will:")
add_numbered([
    "Define the scope of work and what we expect to get back.",
    "Confirm billing rates and any rate caps (see Attachment C).",
    "Set a budget and get approvals if it's over the threshold (see Attachment A).",
    "Send a written engagement letter covering terms, staffing, and reporting.",
])

add_sub_header("Budgets and Spend")
add_para("Outside counsel give us a budget up front. Legal reviews spend against budget every month.")
add_para("Budget rules:")
add_bullets([
    "The responsible Legal attorney reviews and approves invoices each month.",
    "Compare actual spend to budget and explain any variance.",
    "CFO gets a quarterly report of total legal spend.",
    "If spend will exceed the budget by more than 15%, get pre-approval.",
])

add_sub_header("Strategy and Settlement")
add_para(
    "Legal sets litigation strategy with executive leadership. "
    "Settlement authority follows the Delegation of Authority Matrix (Attachment A). "
    "Put every settlement proposal in writing, including its financial and strategic impact."
)

add_sub_header("Reporting on Litigation")
add_para("Legal gives executive leadership regular updates. Each update covers:")
add_bullets([
    "Status of every active matter.",
    "Changes in estimated exposure.",
    "Key developments and upcoming deadlines.",
    "Budget vs. actual spend.",
    "Decisions or actions we recommend.",
])
add_para("Reports pull from the Legal Risk Tracker and the Legal Risk Dashboard (Section 10 and Attachment B).")

add_sub_header("Review After a Matter Closes")
add_para(
    "After a significant matter ends, Legal holds a post-mortem. "
    "We look at outside counsel performance, how actual outcomes compared to our original exposure estimate, "
    "and recurring issues that might mean changing a lease form, operational step, or risk practice."
)

# ============================================================
# SECTION 7: CORPORATE GOVERNANCE & ENTITY MANAGEMENT
# ============================================================
page_break()
add_section_header("7. Corporate Governance and Entity Management")
add_para(
    "Legal oversees the Company's governance and keeps accurate records for every entity in the Company structure. "
    "Legal also makes sure the Company stays qualified as a REIT under the Internal Revenue Code (IRC Sections 856\u2013860), "
    "and coordinates with Finance on SEC reports. "
    "These procedures meet what lenders, joint venture partners, and regulators expect."
)
add_para("What these procedures do:")
add_bullets([
    "Keep the Company qualified as a REIT (IRC Sections 856\u2013860).",
    "Keep entity records accurate in every jurisdiction.",
    "Support financings, deals, and loan covenant compliance.",
    "Preserve institutional knowledge of the legal structure.",
    "Coordinate SEC filings (10-K, 10-Q, 8-K) with Finance.",
    "Give employees a way to report compliance concerns safely.",
])

add_sub_header("A. Entity Management")
add_para(
    "Legal keeps entity records in a dedicated Entity Management System (for example, Diligent Entities or CSC Entity Management). "
    "This system is the single source of truth. Do not use a spreadsheet as the main record."
)
add_para("For every entity, the system tracks at least:")
add_bullets([
    "Legal name, entity type (LLC, LP, corporation, trust, etc.), and EIN.",
    "State of formation and formation date.",
    "Ownership structure, percentages, and org chart.",
    "Governing documents (operating agreement, LP agreement, bylaws, declaration of trust).",
    "Registered agent name, address, and contact info for each state.",
    "Principal office and service of process address.",
    "Good standing status in the home state and every state of qualification.",
    "Officers, managers, directors, and authorized signers with appointment and removal dates.",
    "Material intercompany contracts and loans.",
])
add_para("Legal handles these entity actions:")
add_bullets([
    "Forming and dissolving entities, including all organizational documents and filings.",
    "Amending governing documents (with required approvals under the authority matrix).",
    "Restructurings, mergers, and conversions between entities.",
    "Foreign qualification and withdrawal filings.",
    "Changes to registered agents or registered offices.",
])

add_sub_header("B. Annual Entity Records Review")
add_para("Every year (by the end of Q1), Legal runs a full review of entity records. The steps:")
add_numbered([
    "Pull a full entity register from the Entity Management System: name, formation state, qualification states, registered agent, and good standing.",
    "Cross-check the register against tax records, org charts, and financial statements. Find any gaps or stale data.",
    "For each entity, verify (a) governing documents are current, (b) ownership matches cap tables and tax filings, (c) officer and manager lists are current, and (d) registered agent info is correct.",
    "Write an Entity Records Attestation Report summarizing what we found, any problems, and what's being done about them.",
    "The GC signs the report and sends it to the CFO and Board (or governing body) by April 30 each year.",
])
add_para(
    "Also update entity records right away whenever something material happens: an acquisition, disposition, financing, "
    "refinancing, joint venture, or restructuring."
)

add_sub_header("C. Governance Records")
add_para("Legal keeps governance records organized, including:")
add_bullets([
    "Member, manager, and board approvals and minutes.",
    "Written consents and resolutions (including unanimous written consents instead of meetings).",
    "Officer, manager, and director appointments and removals.",
    "Delegation of authority documents.",
    "Annual conflict-of-interest questionnaires and disclosures.",
    "Related-party transaction approvals.",
])
add_para(
    "Save all governance actions in SharePoint. Track recurring tasks (annual meetings, deliverables, filing deadlines) "
    "on the Governance Calendar in Section 7.H."
)

add_sub_header("D. Authority and Approvals")
add_para("Legal runs the Company's authority framework (see Attachment A).")
add_para("Legal's responsibilities:")
add_bullets([
    "Confirm who can sign each contract, lease, loan document, or deed.",
    "Get approvals for material transactions (including board or member consent).",
    "Keep the authorized signatory list current (and specimen signatures if lenders or title companies require them).",
    "Check approval thresholds before anyone signs.",
    "Coordinate with Acquisitions and Asset Management on authority for acquisitions, dispositions, and financings.",
])
add_para(
    "Legal sends an updated authorized signatory list to Finance, Asset Management, Acquisitions, and Property Management "
    "at least once a year and within 5 business days of any change."
)

add_sub_header("E. REIT Compliance")
add_para(
    "The Company has elected REIT status under the Internal Revenue Code. Keeping that status matters a lot: "
    "losing it means federal corporate tax and could trigger defaults under our credit facilities and JV agreements. "
    "Legal works with Finance and outside tax advisors to track these REIT rules:"
)
add_para("**Asset Tests (IRC 856(c)(4))**")
add_bullets([
    "At the end of each quarter, at least 75% of the Company's assets (by value) must be real estate, cash, cash items, or government securities (the \"75% Asset Test\").",
    "No more than 25% of assets can be other securities.",
    "No more than 5% of assets in any one issuer, and no more than 10% of any issuer's voting or total securities (other than TRSs or qualified REIT subsidiaries) (the \"5%/10% Tests\").",
    "Total taxable REIT subsidiary (TRS) securities can't exceed 20% of assets.",
])
add_para("**Gross Income Tests (IRC 856(c)(2) and (3))**")
add_bullets([
    "At least 75% of gross income each year must come from rent on real property, mortgage interest, sale of real estate, REIT dividends, or other qualifying real estate income (the \"75% Income Test\").",
    "At least 95% of gross income must come from the 75% sources plus dividends, interest, and gain on securities (the \"95% Income Test\").",
    "Legal works with Finance to make sure rents qualify under IRC 856(d): watch for impermissible tenant services, related-party rents, and personal property rent above 15% of total rent.",
])
add_para("**Distribution Requirement (IRC 857(a))**")
add_bullets([
    "The Company must distribute at least 90% of REIT taxable income (before the dividends-paid deduction and excluding net capital gains) each year.",
    "Legal works with Finance to run distribution math within 60 days after year-end and declare and pay the distribution on time.",
    "Legal watches preferential dividend rules and consent dividend procedures as needed.",
])
add_para("**Organizational Tests (IRC 856(a) and (b))**")
add_bullets([
    "The Company must be managed by one or more trustees or directors.",
    "Ownership must be in transferable shares or certificates.",
    "At least 100 shareholders for at least 335 days of every 12-month tax year (the \"100 Shareholder Test\").",
    "During the last half of any tax year, no more than 50% of shares (by value) can be owned, directly or constructively, by 5 or fewer individuals (the \"5/50 Test\").",
])
add_para(
    "Legal maintains a REIT Compliance Checklist and updates it quarterly with Finance and outside tax counsel. "
    "If a REIT rule might have been broken, tell the GC and CFO immediately. Build a fix-it plan within 5 business days."
)

add_sub_header("F. State Filings")
add_para(
    "Legal works with an outside registered agent to keep every Company entity in good standing in every state. "
    "State filings fall into these buckets:"
)
add_para("**Annual reports and statements of information**")
add_bullets([
    "Required by the home state and every state where the entity is qualified.",
    "Deadlines vary. Legal keeps a master filing calendar in the Entity Management System with reminders at 60 and 30 days before each deadline.",
])
add_para("**Franchise and business tax filings**")
add_bullets([
    "Franchise tax returns, gross receipts tax, and similar state-level filings needed for good standing (for example, Delaware franchise tax, Texas franchise tax, California LLC annual tax).",
    "Legal works with Finance and outside tax advisors to file and pay these on time.",
])
add_para("**Foreign qualification and withdrawal**")
add_bullets([
    "Apply for authority to do business in any state where a Company entity owns property, has an office, or otherwise triggers qualification.",
    "File withdrawals when an entity leaves a state.",
    "When the Company buys or sells property in a new state, or starts operating in a new state, Legal checks whether qualification is needed.",
])
add_para("**Registered agent or office changes**")
add_bullets([
    "File the paperwork to update the Secretary of State whenever the registered agent or office changes.",
])
add_para("**If we miss a filing**")
add_numbered([
    "Tell the GC right away.",
    "Within 2 business days: figure out what happened, whether good standing is affected, and what the penalties are.",
    "Within 5 business days: file the late paperwork with any penalties or fees.",
    "If good standing is lost or suspended: start reinstatement immediately. Tell the CFO and any lenders or counterparties that need good standing certificates.",
    "Write an incident report: what went wrong, root cause, how we fixed it, how we'll prevent it. Save it in the Entity Management System and include it in the quarterly compliance review.",
    "The GC tells executive leadership within 5 business days if the miss causes a loss of good standing or material penalty.",
])

add_sub_header("G. SEC Reporting")
add_para("Legal supports SEC reports with Finance, the auditors, and outside securities counsel. Legal's role:")
add_bullets([
    "Review and give legal input on the annual Form 10-K: legal proceedings (Item 3), risk factor updates (Item 1A), and legal contingency footnotes.",
    "Review and give legal input on quarterly Form 10-Qs: updates to legal proceedings and contingencies.",
    "Coordinate Form 8-K filings for reportable events: material definitive agreements, major acquisitions or dispositions, new direct financial obligations, changes in control, and officer or director changes.",
    "Keep a log of possible 8-K trigger events and monitor activity across Acquisitions, Finance, and Asset Management to catch reportable events quickly.",
    "Support proxy statements, information statements, and other SEC filings.",
    "Coordinate Section 16 filings (Forms 3, 4, 5) for directors, officers, and big shareholders with outside securities counsel.",
])
add_para(
    "Legal also sits on the Company's disclosure committee (or equivalent) and keeps an SEC Reporting Calendar "
    "with filing deadlines, earnings release dates, and blackout windows."
)

add_sub_header("H. Governance and Compliance Calendars")
add_para(
    "Legal maintains two calendars in the Entity Management System with automated alerts: "
    "the Governance Calendar and the REIT Compliance Calendar. Review both at the quarterly Finance Coordination meeting (Section 9)."
)
add_para("**Governance Calendar** includes at least:")
add_bullets([
    "Annual meetings (or written consents) for every entity that needs them.",
    "Board and committee meeting dates and deliverable deadlines.",
    "Annual officer and manager appointment or reappointment dates.",
    "Annual conflict-of-interest questionnaire send and return dates.",
    "Annual delegation of authority review and recirculation.",
    "Authorized signatory list update and distribution.",
    "Entity Records Attestation Report deadline (April 30).",
    "Annual insurance renewal coordination with Risk Management.",
    "SEC filing deadlines (10-K, 10-Q, 8-K, proxy).",
])
add_para("**REIT Compliance Calendar** covers at least:")

add_table(
    ["Item", "Frequency", "Who", "Deadline"],
    [
        ["Asset Test review (75%, 5%/10%, TRS 20%)", "Quarterly", "Finance + Legal + Tax Advisor", "Within 30 days after quarter-end"],
        ["Gross Income Test review (75% and 95%)", "Quarterly", "Finance + Legal + Tax Advisor", "Within 30 days after quarter-end"],
        ["Distribution Requirement math", "Annually", "Finance + Legal + Tax Advisor", "Within 60 days after tax year-end"],
        ["100 Shareholder Test", "Annually", "Finance + Transfer Agent", "Before year-end (335-day count)"],
        ["5/50 Ownership Concentration Test", "Semi-annually", "Finance + Transfer Agent", "Second half of each tax year"],
        ["TRS Activity and Income review", "Quarterly", "Finance + Legal + Tax Advisor", "Within 30 days after quarter-end"],
        ["Related-Party Rent review (IRC 856(d))", "Quarterly", "Finance + Legal + Tax Advisor", "Within 30 days after quarter-end"],
        ["Impermissible Tenant Service Income review", "Quarterly", "Finance + Legal + Asset Mgmt", "Within 30 days after quarter-end"],
        ["Prohibited Transaction screen (IRC 857(b)(6))", "Per deal", "Legal + Tax Advisor", "Before every disposition closes"],
        ["Annual REIT Compliance Certification", "Annually", "GC + CFO", "Within 90 days after tax year-end"],
        ["State Filing Compliance review", "Quarterly", "Legal + Registered Agent", "15th of the month after quarter-end"],
        ["Good Standing Certificate collection", "Annually", "Legal + Registered Agent", "By March 31"],
    ],
    [2.0, 1.0, 1.8, 1.7]
)

add_sub_header("I. Reporting Compliance Concerns (Whistleblower)")
add_para(
    "The Company wants employees, officers, contractors, and others to report suspected problems without fear of retaliation. "
    "Legal runs the internal reporting process."
)
add_para("**How to report.** Use any of these channels:")
add_bullets([
    "Directly to the GC or any attorney in Legal (in person, phone, or email).",
    "The Company's anonymous ethics and compliance hotline, run by a third party (phone or web).",
    "In writing to the Chair of the Audit Committee, mailed to the Company's registered office and marked \"Confidential \u2013 Audit Committee.\"",
])
add_para("**What you can report.** The process covers:")
add_bullets([
    "Suspected fraud, embezzlement, or financial problems.",
    "Violations of securities laws, REIT rules, or tax laws.",
    "Violations of the Company's Code of Business Conduct and Ethics.",
    "Suspected environmental, health, or safety violations.",
    "Suspected violations of fair housing laws, the Americans with Disabilities Act (ADA), or other anti-discrimination laws.",
    "Undisclosed conflicts of interest or related-party deals.",
    "Retaliation against anyone who reported in good faith.",
])
add_para("**What happens next**")
add_numbered([
    "Legal (or the hotline provider) logs the report and assigns a tracking number within 1 business day.",
    "The GC does an initial review within 5 business days to figure out what's needed.",
    "Financial reporting, accounting fraud, or audit issues go to the Audit Committee or its designee.",
    "Legal runs the investigation, bringing in outside counsel or forensic specialists if needed.",
    "When the investigation is done, Legal writes up findings and recommended actions.",
    "The GC reports open and recently closed investigations to the Audit Committee at least quarterly.",
])
add_para("**No retaliation.**")
add_para(
    "The Company prohibits retaliation against anyone who reports in good faith or takes part in an investigation. "
    "Retaliators face discipline, up to termination. Legal includes these protections in the employee handbook and covers them in annual training."
)

# ============================================================
# SECTION 8: CONFIDENTIALITY & INFORMATION HANDLING
# ============================================================
page_break()
add_section_header("8. Confidentiality and Information Handling")
add_para(
    "Legal handles sensitive information: privileged communications, litigation materials, business secrets, tenant personal data, "
    "and compliance records. Handling this properly protects the Company and keeps us compliant with privacy, fair housing, and anti-discrimination laws."
)

add_sub_header("A. How We Classify Documents")
add_para(
    "Every document Legal creates, receives, or keeps gets one of four classifications. "
    "The classification tells you how to handle, store, share, and destroy it."
)

add_table(
    ["Level", "What It Means", "How to Handle It"],
    [
        ["PRIVILEGED",
         "Attorney-client communications and attorney work product.",
         "Only Legal and specifically authorized people can see it. Must carry the privilege header. "
         "Cannot leave the Company without GC approval. Store in access-controlled SharePoint folders with audit logs on."],
        ["CONFIDENTIAL",
         "Not privileged but commercially sensitive: deal terms, tenant financials, settlement amounts, REIT data, personnel matters, whistleblower reports.",
         "Only share with people who need it for their job. Label \"CONFIDENTIAL.\" Store in access-controlled SharePoint folders. "
         "Cannot leave the Company without GC or authorized officer approval."],
        ["INTERNAL",
         "For internal use, not public, but not Confidential: internal policies, standard forms, general legal updates.",
         "Share freely inside the Company. Store in SharePoint. Don't share outside without department-head approval."],
        ["PUBLIC",
         "Approved for public release: SEC filings, press releases, recorded documents.",
         "No sharing restrictions. Store in SharePoint and public filing systems."],
    ],
    [1.3, 2.5, 2.7]
)

add_para(
    "Pick the right level when you create the document. If you're not sure, treat it as CONFIDENTIAL until the GC or a designated attorney reviews it."
)

add_sub_header("B. What Counts as Confidential")
add_para("Confidential information includes (but isn't limited to):")
add_bullets([
    "Legal advice and attorney-client communications (these are PRIVILEGED).",
    "Litigation strategy, work product, and settlement talks.",
    "Tenant disputes, negotiations, and financial information.",
    "Regulatory matters and government inquiries.",
    "Deal information (purchase, sale, due diligence, financing).",
    "REIT compliance data (asset tests, income tests, distributions).",
    "Personnel matters (including whistleblower reports and investigation files).",
    "Fair housing complaints, ADA accommodation requests, and anti-discrimination investigations.",
    "Tenant personal information (financials, tax returns, guarantor info).",
])

add_sub_header("C. Attorney-Client Privilege")
add_para(
    "Communications with Legal that are meant to give or ask for legal advice are usually privileged. "
    "Keeping privilege protects the Company's ability to get candid advice and defend itself."
)
add_para("Everyone should:")
add_bullets([
    "Only share privileged messages with people who need the advice.",
    "Don't forward privileged emails outside the Company, or to anyone inside who wasn't on the original, without Legal's approval.",
    "Don't copy or summarize legal advice into non-privileged documents (like board decks or memos to non-lawyers) without checking with Legal first.",
    "Label privileged documents and emails with the privilege header.",
    "Don't discuss legal advice in front of third parties (consultants, vendors, JV partners) unless Legal confirms a common-interest or joint-defense privilege applies.",
])
add_para(
    "Privilege header to use: \"PRIVILEGED AND CONFIDENTIAL \u2013 ATTORNEY-CLIENT COMMUNICATION. "
    "This communication is protected by the attorney-client privilege and/or the work product doctrine. "
    "Do not forward, copy, or disclose without prior authorization of the Legal Department.\""
)
add_para("Legal runs privilege training for all employees at least once a year and gives new hires written guidance during onboarding.")

add_sub_header("D. Digital Communications and Security")
add_para("Rules for handling privileged or confidential information in email, chat, video, cloud storage, and mobile devices:")
add_numbered([
    "**Email.** Use only the Company's approved email system. Never send or receive privileged or confidential information from a personal account. Put the privilege header on privileged emails. Put \"CONFIDENTIAL\" in the subject line for confidential emails.",
    "**Chat and collaboration platforms.** Don't send privileged legal advice over chat (Teams, Slack) unless Legal and IT have approved the platform for it and the platform keeps messages per our retention policy. Keep confidential info in access-restricted channels.",
    "**Video conferencing.** Check who's on the call. Use password-protected meetings. Don't record unless Legal has approved recording for that meeting. Make sure no one in the room can overhear who shouldn't.",
    "**Cloud storage and file sharing.** Store privileged and confidential documents only in approved systems (SharePoint, ShareFile). Never use personal Dropbox, Google Drive, or unapproved file-sharing services. Set shared links to expire and to require Company login.",
    "**Mobile devices.** Follow the Company's mobile device management (MDM) policy: encryption, passcode, and remote wipe.",
])

add_sub_header("E. Document Storage and Retention")
add_para(
    "Legal documents live in SharePoint or ShareFile (per the schedule below). Never store them on your laptop, personal device, or an unapproved platform. "
    "Only people with a business need get access, and Legal reviews access permissions quarterly."
)
add_para("Retention schedule:")

add_table(
    ["Document Type", "How Long to Keep", "Where to Store"],
    [
        ["Executed Leases", "Life of lease + 7 years", "ShareFile"],
        ["Lease Drafts and Negotiations", "Life of lease + 3 years", "SharePoint"],
        ["Litigation Files", "Resolution + 7 years", "SharePoint / Outside Counsel"],
        ["Corporate Governance Records", "Permanent", "SharePoint / Entity Mgmt System"],
        ["REIT Compliance Records", "Permanent", "SharePoint"],
        ["SEC Filings and Support", "Permanent", "SharePoint"],
        ["Outside Counsel Invoices", "7 years", "SharePoint / Finance"],
        ["Fair Housing / ADA Complaint Files", "Resolution + 7 years", "SharePoint (restricted)"],
        ["Whistleblower Reports and Investigations", "Resolution + 7 years", "SharePoint (restricted)"],
        ["General Correspondence", "3 years", "SharePoint"],
        ["Regulatory Filings", "Permanent", "SharePoint"],
        ["Entity Formation and Qualification Documents", "Life of entity + 7 years", "Entity Mgmt System / SharePoint"],
    ],
    [2.2, 2.2, 2.1]
)

add_sub_header("F. Fair Housing, ADA, and Anti-Discrimination")
add_para(
    "The Company's properties and operations have to follow federal, state, and local fair housing, accessibility, and anti-discrimination laws. "
    "That includes the Fair Housing Act (42 U.S.C. \u00a7\u00a7 3601\u20133619), the ADA (42 U.S.C. \u00a7\u00a7 12101 et seq.), Section 504 of the Rehabilitation Act, "
    "and state and local human rights laws. Legal's role:"
)
add_bullets([
    "Advise Property Management and Asset Management on fair housing and ADA rules, including how to handle reasonable accommodation and reasonable modification requests.",
    "Review and approve fair housing policies, marketing materials, and tenant selection criteria.",
    "Coordinate the Company's response to fair housing complaints, HUD complaints, and state or local human rights commission inquiries.",
    "Keep a log (classified CONFIDENTIAL) of fair housing and ADA complaints, accommodation requests, and investigations.",
    "Support the Company's ADA program, including working with architects, engineers, and accessibility consultants on ADA surveys, barrier removal, and accessibility in new construction or renovations.",
    "Store all fair housing and ADA records per Section 8.E in access-restricted SharePoint folders.",
    "Cover fair housing and anti-discrimination topics in the annual training program (Section 10).",
])

add_sub_header("G. Litigation Holds")
add_para(
    "As soon as a matter could lead to litigation, arbitration, a regulatory investigation, or a government inquiry, Legal puts a litigation hold in place. "
    "A hold means preserving all potentially relevant documents and communications. If we don't preserve them, we can face sanctions or adverse inferences."
)
add_para("Litigation hold steps:")
add_numbered([
    "Legal sends a written hold notice to every relevant custodian. It names the matter, lists the document categories to preserve, and tells them to stop any routine deletion or destruction of those documents.",
    "Each custodian acknowledges the hold in writing (email is fine) within 3 business days. Legal follows up with anyone who hasn't acknowledged.",
    "Legal works with IT to stop any automatic deletion, archiving, or overwriting of documents, emails, chat messages, and other electronically stored info covered by the hold.",
    "Legal sends hold reminders at least every 90 days while the hold is active.",
    "When the matter ends, Legal sends a written release. Custodians hear that the hold is lifted and regular retention schedules can resume.",
    "Legal keeps a litigation hold log: issuance date, custodians notified, acknowledgments, and release date.",
])

add_sub_header("H. Talking to Outsiders")
add_para("Only people authorized by the GC can speak for the Company to:")
add_bullets([
    "Outside counsel.",
    "Opposing counsel or opposing parties.",
    "Regulators, government agencies, and law enforcement.",
    "Fair housing agencies, human rights commissions, and HUD.",
    "The SEC, state securities regulators, and stock exchange representatives.",
])
add_para(
    "Route all external legal communications through Legal. Never answer a subpoena, civil investigative demand, regulatory inquiry, "
    "or government information request without checking with Legal first. Legal keeps a log of material external legal communications."
)

# ============================================================
# SECTION 9: COMMUNICATION & ESCALATION PROTOCOLS
# ============================================================
page_break()
add_section_header("9. Communication and Escalation")
add_para(
    "Clear communication and fast escalation keep legal risk manageable. "
    "This section says when and how to escalate to Legal and executive leadership."
)

add_sub_header("A. Regular Meetings")
add_para("Legal works with:")
add_bullets([
    "Asset Management (leasing and tenants).",
    "Property Management (operations and incidents).",
    "Finance (reserves, deals, governance).",
])
add_para("Standing meetings:")

add_table(
    ["Meeting", "Participants", "How Often", "Purpose"],
    [
        ["Leasing Pipeline Review", "Legal + Asset Mgmt", "Weekly", "Pending lease requests and priorities"],
        ["Litigation/Risk Update", "Legal + Executive Leadership", "Bi-weekly", "Active matters, exposure, strategic calls"],
        ["Finance Coordination", "Legal + Finance/CFO", "Monthly", "Reserves, legal spend, governance filings"],
        ["Property Ops Sync", "Legal + Property Mgmt", "As needed", "Incidents, claims, vendor disputes"],
    ],
    [1.5, 1.5, 1.2, 2.3]
)

add_sub_header("B. When to Escalate to Legal")
add_para("Escalate to Legal anything that involves:")
add_bullets([
    "Possible or actual litigation.",
    "Tenant bankruptcy or insolvency.",
    "Big financial exposure.",
    "Regulatory or government inquiries.",
    "Serious personal injury or property damage.",
    "Disputes affecting Company rights or obligations.",
])

add_sub_header("C. How Fast to Escalate")
add_para("Escalation timelines:")

add_table(
    ["Event", "Timeline", "Who to Tell"],
    [
        ["Lawsuit or demand letter arrives", "Same business day", "General Counsel"],
        ["Regulatory inquiry or subpoena", "Same business day", "General Counsel + CEO"],
        ["Tenant files bankruptcy", "Within 24 hours", "General Counsel + Asset Mgmt Lead"],
        ["Serious personal injury on property", "Immediately", "General Counsel + Property Mgmt Lead"],
        ["Financial exposure over $250,000", "Within 24 hours", "General Counsel + CFO"],
        ["Possible privilege waiver or breach", "Immediately", "General Counsel"],
        ["Insurance coverage dispute", "Within 48 hours", "General Counsel + Asset Mgmt Lead"],
    ],
    [2.2, 2.2, 2.1]
)

add_sub_header("D. When Legal Escalates to Executives")
add_para("Legal escalates to executive leadership when a matter involves:")
add_bullets([
    "Big financial exposure.",
    "Reputational risk.",
    "Impact across the portfolio.",
    "Strategic impact on the business.",
])
add_para(
    "Escalations to executives include a written summary, current status, recommended action, and financial impact. "
    "Legal uses a standard escalation memo format for consistency."
)

add_sub_header("E. Regular Legal Updates")
add_para("Legal gives executive leadership regular updates with:")
add_bullets([
    "Status of significant matters.",
    "Litigation exposure.",
    "Key risks and developments.",
])
add_para("The Legal Risk Tracker and Legal Risk Dashboard (Section 10 and Attachment B) feed these updates.")

# ============================================================
# SECTION 10: LEGAL RISK TRACKING & REPORTING
# ============================================================
page_break()
add_section_header("10. Legal Risk Tracking and Reporting")
add_para("Legal runs a Legal Risk Tracker so everyone can see what's going on across the portfolio.")
add_para("The tracker lets us:")
add_bullets([
    "Monitor lawsuits and disputes.",
    "Track tenant defaults and exposure.",
    "Spot trends and recurring issues.",
    "Support executive decisions.",
])

add_sub_header("A. Risk Categories")
add_para("Legal sorts matters into:")
add_bullets([
    "Litigation.",
    "Tenant defaults.",
    "Bankruptcy.",
    "Insurance claims.",
    "Regulatory matters.",
    "Environmental issues.",
])

add_sub_header("B. How We Rate Risk")
add_para("Each matter gets a probability-impact rating:")

add_table(
    ["Rating", "Probability", "Financial Impact"],
    [
        ["High", "Likely (>70%)", "> $500,000"],
        ["Medium", "Possible (30\u201370%)", "$100,000 \u2013 $500,000"],
        ["Low", "Unlikely (<30%)", "< $100,000"],
    ],
    [2.2, 2.2, 2.1]
)

add_para("We also consider:")
add_bullets([
    "Strategic or precedent risk.",
    "Reputational impact.",
    "Whether insurance covers it.",
    "Strength of our legal position.",
])
add_para("Review ratings every reporting cycle. Document why a rating changes.")

add_sub_header("C. Reporting Cadence")
add_para("How often we update the tracker:")
add_bullets([
    "High-priority matters: weekly.",
    "Medium-priority matters: monthly.",
    "Low-priority matters: quarterly.",
])
add_para(
    "Send updates to executive leadership before the scheduled legal update meetings. "
    "Prepare the Legal Risk Dashboard (Attachment B) monthly for the CFO and quarterly for the Board."
)

add_sub_header("D. Reserves")
add_para(
    "Finance uses the risk data to set reserves. Legal gives estimated exposure ranges for every Medium or High matter, "
    "updated at least quarterly or right away if something material changes."
)

add_sub_header("E. What We Do with Risk Data")
add_para("Risk data helps:")
add_bullets([
    "Set reserves.",
    "Support strategic planning.",
    "Identify where to reduce risk.",
    "Improve lease forms and procedures.",
])

# ============================================================
# SECTION 11: TRAINING & CONTINUOUS IMPROVEMENT
# ============================================================
page_break()
add_section_header("11. Training and Continuous Improvement")
add_para("Legal runs training to help other teams work with us and reduce risk.")

add_sub_header("A. Internal Training Program")
add_para("Training schedule:")

add_table(
    ["Audience", "How Often", "Topics"],
    [
        ["Asset Management", "Quarterly", "Lease provisions, negotiation, document workflow"],
        ["Property Management", "Quarterly", "Incident reporting, insurance claims, escalation"],
        ["Finance", "Twice a year", "Risk reporting, reserves, governance"],
        ["All Departments", "Annually", "Confidentiality, privilege, policy updates"],
        ["Legal new hires", "At onboarding", "Systems, standard forms, naming, workflows"],
    ],
    [1.8, 1.5, 3.2]
)

add_sub_header("B. New Hire Onboarding (Legal)")
add_para("New Legal staff complete this checklist:")
add_numbered([
    "Read the Manual and all attachments.",
    "Learn the Company's entity structure and portfolio.",
    "Train on Salesforce (intake, workflow, dashboard).",
    "Train on SharePoint (storage, collaboration, naming).",
    "Review standard lease forms and key provisions.",
    "Review active litigation and the risk tracker.",
    "Meet the outside counsel panel and learn billing guidelines.",
    "Review confidentiality and privilege rules.",
])

add_sub_header("C. Training Materials")
add_para(
    "Decks, reference guides, and recorded sessions live in a dedicated SharePoint folder Legal maintains. "
    "Legal updates them whenever procedures, systems, or forms change."
)

add_sub_header("D. Legal Updates")
add_para(
    "Legal watches for legal and market developments and shares updates with internal teams. "
    "Updates go out by email memo, training sessions, or at standing interdepartmental meetings."
)

add_sub_header("E. Continuous Improvement")
add_para("Legal regularly reviews:")
add_bullets(["Leasing processes.", "Litigation outcomes.", "Recurring issues."])
add_para("To find chances to:")
add_bullets(["Improve standard lease forms.", "Streamline workflows.", "Reduce legal exposure."])
add_para(
    "After big matters close, Legal runs a post-mortem to capture lessons. "
    "We work those lessons into updated procedures, training, and standard forms."
)

# ============================================================
# SECTION 12: TECHNOLOGY & WORKFLOW OPTIMIZATION
# ============================================================
page_break()
add_section_header("12. Technology and Workflow")
add_para("Legal uses technology to work faster, stay organized, and scale as the Company grows.")

add_sub_header("A. Systems We Use")
add_para("Current tools:")

add_table(
    ["System", "What It Does", "Owner", "Users"],
    [
        ["Salesforce", "Lease intake, workflow, pipeline dashboard", "Legal / Asset Mgmt", "Legal, Asset Mgmt, Executive"],
        ["SharePoint", "Document collaboration, version control, storage", "Legal / IT", "Legal, Asset Mgmt, Finance"],
        ["ShareFile", "Executed document storage and external sharing", "Asset Mgmt", "Legal, Asset Mgmt"],
        ["Legal Risk Tracker", "Litigation and risk monitoring, reporting", "Legal", "Legal, Executive, CFO"],
    ],
    [1.3, 2.5, 1.3, 1.4]
)

add_para("Access is role-based. Legal works with IT to keep access controls right.")

add_sub_header("B. System Integration")
add_para("Legal aims to:")
add_bullets([
    "Tie legal workflows to operational systems.",
    "Capture data consistently across tools.",
    "Make information available to the people who need it.",
    "Cut duplicate data entry where systems can integrate.",
])
add_para(
    "Legal works with IT and Finance so legal data (lease milestones, litigation reserves, entity records) "
    "shows up accurately in the Company's financial and operational systems."
)

add_sub_header("C. Picking New Tools")
add_para("Legal may look at new tools to improve:")
add_bullets([
    "Lease data extraction and analysis.",
    "Document automation.",
    "Risk tracking and reporting.",
    "Workflow automation.",
])
add_para("When evaluating, consider:")
add_bullets([
    "Can it integrate with Salesforce and SharePoint?",
    "Does it meet our data security and access standards?",
    "Does the benefit justify the cost?",
    "How hard is it to adopt and train on?",
    "Is the vendor reliable and well-supported?",
])
add_para(
    "Coordinate with IT and Finance before buying or rolling out anything new. "
    "Material tech purchases go through the Company's standard procurement process."
)

add_sub_header("D. Keeping Data Clean")
add_para("Keeping data consistent across systems matters. Legal:")
add_bullets([
    "Audits data in Salesforce and SharePoint periodically.",
    "Makes sure lease milestones, entity records, and litigation data stay current.",
    "Reports data problems to the system owner for fixing.",
    "Keeps naming and folder structures consistent everywhere.",
])

add_sub_header("E. Change Management")
add_para(
    "If a system upgrade, migration, or workflow change will affect Legal, coordinate with Legal first. "
    "Legal helps test and train before the change goes live so legal operations don't get disrupted."
)

# ============================================================
# SECTION 13: POLICY REVIEW AND UPDATES
# ============================================================
page_break()
add_section_header("13. Updating This Manual")
add_para("This Manual is a living document. Expect updates as operations and laws change.")

add_sub_header("A. Annual Review")
add_para("Legal reviews the Manual once a year, on the Company's fiscal year. The review checks for alignment with:")
add_bullets([
    "Company structure or strategy changes.",
    "New systems or upgrades.",
    "Regulatory developments.",
    "Operational needs.",
])
add_para("Finish the annual review in Q1 of each fiscal year.")

add_sub_header("B. Updates")
add_para(
    "Legal updates the Manual as needed. The GC has to approve material changes, and Legal tells affected teams."
)
add_para("Version control rules:")
add_bullets([
    "Every revision has a version number and effective date.",
    "A change log tracks all material updates (below).",
    "The current version lives in a designated SharePoint location available to Legal staff.",
    "Prior versions are archived for reference.",
])

add_sub_header("C. Change Log")

add_table(
    ["Version", "Date", "Change", "Approved By"],
    [
        ["1.0", "[Date]", "Initial publication", "[Name/Title]"],
        ["", "", "", ""],
        ["", "", "", ""],
    ],
    [0.9, 1.1, 3.3, 1.2]
)

add_sub_header("D. Implementation")
add_para(
    "Legal implements and enforces the procedures in this Manual. "
    "All Legal staff are expected to read the current version and follow it."
)

# ============================================================
# ATTACHMENT A: DELEGATION OF AUTHORITY MATRIX
# ============================================================
page_break()
add_section_header("Attachment A: Who Can Approve What")
add_para(
    "This matrix shows who can approve each type of transaction. "
    "Executive leadership sets the dollar thresholds. Anything above the top tier needs Board approval."
)

add_table(
    ["Action / Document", "Tier 1: GC", "Tier 2: GC + CFO", "Tier 3: GC + CEO", "Tier 4: Board"],
    [
        ["New Lease Execution", "< $[X] annual rent", "$[X] \u2013 $[Y]", "$[Y] \u2013 $[Z]", "> $[Z]"],
        ["Lease Amendment", "< $[X] impact", "$[X] \u2013 $[Y]", "$[Y] \u2013 $[Z]", "> $[Z]"],
        ["Lease Termination / Surrender", "< $[X]", "$[X] \u2013 $[Y]", "$[Y] \u2013 $[Z]", "> $[Z]"],
        ["Litigation Settlement", "< $[X]", "$[X] \u2013 $[Y]", "$[Y] \u2013 $[Z]", "> $[Z]"],
        ["Outside Counsel Engagement", "< $[X] budget", "$[X] \u2013 $[Y]", "> $[Y]", "N/A"],
        ["Vendor / Service Contracts", "< $[X]", "$[X] \u2013 $[Y]", "$[Y] \u2013 $[Z]", "> $[Z]"],
        ["Entity Formation / Dissolution", "GC approval", "GC + CFO", "GC + CEO", "Board"],
        ["Acquisition / Disposition", "N/A", "N/A", "GC + CEO", "Board"],
        ["Financing / Loan Documents", "N/A", "GC + CFO", "GC + CEO", "Board"],
        ["Insurance Claim Settlement", "< $[X]", "$[X] \u2013 $[Y]", "> $[Y]", "N/A"],
    ],
    [1.8, 1.3, 1.3, 1.3, 0.8]
)

add_para("**Notes:**")
add_bullets([
    "$[X], $[Y], $[Z] are dollar thresholds. Executive leadership fills these in before rolling out the matrix.",
    "All approvals must be in writing. Email works for Tier 1.",
    "If the GC is out, the CEO can use Tier 1 and Tier 2 authority for urgent matters. Document it afterward.",
    "Review this matrix every year when we review the Manual (Section 13).",
])

# ============================================================
# ATTACHMENT B: LEGAL RISK DASHBOARD
# ============================================================
page_break()
add_section_header("Attachment B: Legal Risk Dashboard")
add_para(
    "The Legal Risk Dashboard is a one-look summary for executive leadership and the Board. "
    "Legal prepares it monthly and sends it before the scheduled legal update meetings."
)

add_para("**Section 1: Active Matters Summary**")
add_table(
    ["Category", "High Risk", "Medium Risk", "Low Risk", "Total"],
    [
        ["Litigation", "[#]", "[#]", "[#]", "[#]"],
        ["Tenant Defaults", "[#]", "[#]", "[#]", "[#]"],
        ["Bankruptcy", "[#]", "[#]", "[#]", "[#]"],
        ["Insurance Claims", "[#]", "[#]", "[#]", "[#]"],
        ["Regulatory", "[#]", "[#]", "[#]", "[#]"],
        ["Environmental", "[#]", "[#]", "[#]", "[#]"],
        ["TOTAL", "[#]", "[#]", "[#]", "[#]"],
    ],
    [1.8, 1.3, 1.3, 1.1, 1.0]
)

add_para("**Section 2: Financial Exposure Summary**")
add_table(
    ["Risk Level", "Number of Matters", "Estimated Exposure Range", "Current Reserves"],
    [
        ["High", "[#]", "$[X] \u2013 $[Y]", "$[amount]"],
        ["Medium", "[#]", "$[X] \u2013 $[Y]", "$[amount]"],
        ["Low", "[#]", "$[X] \u2013 $[Y]", "$[amount]"],
        ["TOTAL", "[#]", "$[X] \u2013 $[Y]", "$[amount]"],
    ],
    [1.6, 1.6, 1.7, 1.6]
)

add_para("**Section 3: Key Developments**")
add_para("A short narrative covering what happened this period, including:")
add_bullets([
    "New matters opened.",
    "Matters resolved or settled.",
    "Big changes in risk rating or exposure.",
    "Upcoming deadlines or decisions needing executive attention.",
    "Trends or recurring issues worth watching.",
])

add_para("**Section 4: Legal Spend Summary**")
add_table(
    ["Category", "YTD Actual", "YTD Budget", "Variance", "Forecast"],
    [
        ["Litigation", "$[X]", "$[X]", "$[X]", "$[X]"],
        ["Transactional", "$[X]", "$[X]", "$[X]", "$[X]"],
        ["Regulatory/Compliance", "$[X]", "$[X]", "$[X]", "$[X]"],
        ["TOTAL", "$[X]", "$[X]", "$[X]", "$[X]"],
    ],
    [1.8, 1.3, 1.3, 1.1, 1.0]
)

add_para("**Section 5: Action Items**")
add_para("Decisions or approvals the executives or Board need to make, with recommended actions and deadlines.")

# ============================================================
# ATTACHMENT C: OUTSIDE COUNSEL BILLING GUIDELINES
# ============================================================
page_break()
add_section_header("Attachment C: Outside Counsel Billing Guidelines")
add_para(
    "These rules apply to every outside law firm PIR Industrial LLC hires. "
    "Firms must read and follow these as a condition of engagement. Legal can update these rules at any time."
)

add_sub_header("1. Billing Rates")
add_bullets([
    "Put the rate schedule in writing before any work starts.",
    "Any annual rate increase needs GC approval up front.",
    "Maximum hourly rates: Partner $[X]/hr; Senior Associate $[X]/hr; Associate $[X]/hr; Paralegal $[X]/hr.",
    "Blended and alternative fee arrangements are welcome but need GC approval.",
])

add_sub_header("2. Staffing")
add_bullets([
    "Staff lean. Match the team to the complexity of the matter.",
    "Name the responsible partner in the engagement letter. Don't change them without prior approval.",
    "Delegate to the lowest-cost attorney who can do the work well.",
    "No more than two attorneys at a deposition, hearing, or meeting without prior approval.",
    "Summer associates and first-year associates can't bill the Company without prior approval.",
])

add_sub_header("3. Invoices")
add_bullets([
    "Submit invoices monthly, within 45 days of the end of the billing period.",
    "Every invoice must include: matter name and number, work description by timekeeper, hours, rate, and total fees.",
    "Time entries in tenths of an hour (0.1) or finer. No longer blocks.",
    "No block billing (multiple tasks in one time entry).",
    "No vague entries like \"research,\" \"review documents,\" or \"attend to matter.\" Say what you did.",
])

add_sub_header("4. Charges We Won't Pay")
add_para("We won't reimburse:")
add_bullets([
    "Firm overhead (word processing, secretarial, filing, internal copying).",
    "Legal research subscription fees (Westlaw, LexisNexis, etc.).",
    "Travel time (unless pre-approved in writing for out-of-jurisdiction work).",
    "First-class or business-class airfare.",
    "Alcohol.",
    "Internal firm meetings, training, or professional development.",
    "Rush charges for printing, copying, or courier we didn't ask for.",
])

add_sub_header("5. Expenses That Need Pre-Approval")
add_para("Get written pre-approval from Legal for:")
add_bullets([
    "Experts and consultants.",
    "Travel and lodging.",
    "Outsourced discovery or document review.",
    "Filing fees over $1,000.",
    "Any single expense over $2,500.",
])

add_sub_header("6. Matter Budgets")
add_bullets([
    "Give us a matter budget at the start of the engagement, broken down by phase.",
    "Legal has to approve the budget before substantive work begins.",
    "If actual spend will exceed budget by more than 15%, tell Legal in writing and get approval for a revised budget before you keep going.",
    "Submit a quarterly budget-to-actual report with each invoice.",
])

add_sub_header("7. Reporting")
add_para("Outside counsel must provide:")
add_bullets([
    "Monthly status reports: activity, developments, next steps.",
    "Prompt notice of any material development (bad ruling, settlement demand, deadline change).",
    "Quarterly budget-to-actual analysis.",
    "Early case assessment within 30 days of engagement: estimated exposure range, recommended strategy, and timeline.",
])

add_sub_header("8. Invoice Review and Payment")
add_bullets([
    "Legal reviews every invoice and can adjust it.",
    "We can reduce or reject charges that don't follow these rules.",
    "Payment terms: net 60 from receipt of a compliant invoice.",
    "Invoice disputes get worked out between the GC and the responsible partner.",
])

add_sub_header("9. Conflicts and Confidentiality")
add_bullets([
    "Disclose any actual or potential conflicts before engagement.",
    "Treat all Company information as confidential. No disclosure without written authorization.",
    "No press releases, articles, or public statements about Company matters without prior written consent.",
])

# ============================================================
# SAVE
# ============================================================
output_path = os.path.expanduser("~/Downloads/Legal Department Operations Manual - Paralegal Edition.docx")
doc.save(output_path)
# Also save a copy to the worktree for easy access
worktree_copy = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Legal Department Operations Manual - Paralegal Edition.docx",
)
doc.save(worktree_copy)
print(f"Document saved: {output_path}")
print(f"Worktree copy:  {worktree_copy}")

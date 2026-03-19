#!/usr/bin/env python3
"""Generate complete Legal Department Operations Manual for PIR Industrial LLC."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FONT = "Times New Roman"
FONT_SIZE = Pt(12)

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# -- Default style --
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = FONT_SIZE
style.paragraph_format.space_after = Pt(6)

# -- Header/Footer --
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = hp.add_run("PIR Industrial LLC \u2013 Legal Department Operations Manual")
run.font.name = FONT
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("Page ")
run.font.name = FONT
run.font.size = Pt(9)
# Add page number field
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
run2 = fp.add_run()
run2._r.append(fldChar1)
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = ' PAGE '
run3 = fp.add_run()
run3._r.append(instrText)
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run4 = fp.add_run()
run4._r.append(fldChar2)


def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = FONT
    run.font.size = FONT_SIZE
    return p


def add_sub_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.name = FONT
    run.font.size = FONT_SIZE
    return p


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    # Parse **bold** markers
    if "**" in text:
        parts = text.split("**")
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            run.font.name = FONT
            run.font.size = FONT_SIZE
            if i % 2 == 1:
                run.bold = True
            if italic:
                run.italic = True
    else:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = FONT_SIZE
        if bold:
            run.bold = True
        if italic:
            run.italic = True
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        run = p.add_run(item)
        run.font.name = FONT
        run.font.size = FONT_SIZE


def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.clear()
        run = p.add_run(item)
        run.font.name = FONT
        run.font.size = FONT_SIZE


def add_table(headers, rows, col_widths_inches=None):
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = FONT
            run.font.size = Pt(10)

    # Set column widths
    if col_widths_inches:
        for row in table.rows:
            for i, width in enumerate(col_widths_inches):
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # spacing after table
    return table


def page_break():
    doc.add_page_break()


# ============================================================
# SECTION 1: PURPOSE AND SCOPE
# ============================================================
add_section_header("1. Purpose and Scope")
add_para(
    "The purpose of this Legal Department Operations Manual (the \u201CManual\u201D) is to establish a consistent "
    "and scalable framework for the management of legal matters across PIR Industrial LLC (the \u201CCompany\u201D). "
    "This Manual is intended to provide guidance to the Legal Department (the \u201CDepartment\u201D) personnel on "
    "such activities and responsibilities, particularly in cooperation with other internal and external stakeholders, "
    "namely, the Company\u2019s ownership, executive leadership, and personnel in other Company departments."
)
add_para(
    "This Manual will define the Department\u2019s oversight and responsibilities necessary to support informed "
    "business decision-making."
)
add_para(
    "As the Company continues to evolve, implement new technologies, and adjust internal systems and workflows, "
    "the Department may update these policies and procedures to ensure continued effectiveness and alignment with "
    "organizational objectives."
)

# ============================================================
# SECTION 2: OPERATING PRINCIPLES
# ============================================================
page_break()
add_section_header("2. Operating Principles")
add_para(
    "The Department operates in support of the Company\u2019s strategic, operational, and financial objectives. "
    "The principles set forth below guide the Department\u2019s approach to delivering legal services, managing risk, "
    "and collaborating with stakeholders."
)
add_para(
    "These principles are intended to ensure that the Department provides timely, practical, and consistent legal "
    "support while maintaining disciplined risk management and operational efficiency."
)

add_sub_header("Business Alignment")
add_para(
    "The Legal Department supports the Company\u2019s business strategy by providing practical legal guidance that "
    "facilitates (and does not disrupt or interfere with) operations and transactions, while protecting the Company\u2019s "
    "legal and financial interests."
)
add_para(
    "Legal analysis and recommendations should be grounded in an understanding of the Company\u2019s business priorities, "
    "portfolio strategy, and operational realities. The Department strives to balance risk management with the need for "
    "efficient execution of business initiatives."
)

add_sub_header("Speed with Discipline")
add_para(
    "Leasing and operational matters often require prompt legal support in order to maintain transaction momentum and "
    "support tenant relationships."
)
add_para(
    "The Department should deliver timely responses while maintaining appropriate legal review standards, including "
    "document accuracy, version control, and internal coordination. Efficient processes and clearly defined workflows "
    "help ensure that speed does not compromise quality or risk management."
)

add_sub_header("Consistency and Standardization")
add_para(
    "Standardized procedures promote efficiency, clarity, and institutional knowledge. Standardization helps ensure "
    "that legal work product remains reliable and that internal teams can easily navigate legal workflows."
)

add_sub_header("Transparency and Risk Visibility")
add_para(
    "The Legal Department maintains structured reporting tools and communication protocols to ensure that material "
    "legal matters are visible to leadership."
)
add_para(
    "Legal risks should be identified early, documented appropriately, and communicated to relevant stakeholders in "
    "a timely manner. Clear reporting mechanisms help support informed decision-making by executive leadership and "
    "facilitate coordinated responses to legal issues affecting the Company."
)

# FIXED: Centralized Document Management (was duplicate of Transparency section)
add_sub_header("Centralized Document Management")
add_para(
    "All legal documents, including lease drafts, litigation files, governance records, and outside counsel "
    "correspondence, must be stored in the Company\u2019s approved document management systems (currently SharePoint "
    "for active collaboration and ShareFile for executed documents)."
)
add_para(
    "Centralized document management ensures version control integrity, facilitates institutional knowledge transfer, "
    "and supports audit readiness. Documents should not be stored locally or in unapproved platforms. The Department "
    "is responsible for maintaining organized folder structures and enforcing naming conventions across all legal workstreams."
)

add_sub_header("Effective Collaboration")
add_para(
    "The Legal Department works in close coordination with Asset Management, Finance, Property Management, "
    "and executive leadership."
)
add_para(
    "Successful collaboration requires clear communication, shared expectations regarding timelines and responsibilities, "
    "and mutual understanding of each department\u2019s role in the Company\u2019s operations."
)
add_para(
    "Legal supports internal teams by providing guidance, documentation, and risk analysis while respecting the "
    "operational ownership of business decisions by the relevant departments."
)

add_sub_header("Prudent Use of Outside Counsel")
add_para(
    "Outside counsel are engaged to provide specialized expertise, jurisdictional coverage, or additional capacity "
    "when required."
)
add_para(
    "The Department is responsible for selecting and supervising outside counsel, defining engagement scope, monitoring "
    "performance, and managing legal spend. Outside counsel should operate as an extension of the Department and adhere "
    "to the Company\u2019s expectations regarding efficiency, communication, and billing practices."
)

add_sub_header("Continuous Improvement")
add_para(
    "The Department seeks to continually improve its processes through the adoption of efficient workflows, technology "
    "tools, and standardized practices."
)
add_para(
    "This includes evaluating technologies that enhance document management, legal reporting, and access to lease and "
    "portfolio data, as well as maintaining internal training and knowledge resources that support consistent legal operations."
)
add_para("Continuous improvement helps ensure that the Department can stay competitive as the Company evolves.")

add_sub_header("Accountability")
add_para(
    "Members of the Department are expected to maintain high standards of professional judgment, integrity, and "
    "accountability in the performance of their responsibilities."
)
add_para(
    "Legal professionals must exercise sound judgment, safeguard confidential information, and ensure that the "
    "Company\u2019s legal obligations are fulfilled in accordance with applicable laws and internal policies."
)

# ============================================================
# SECTION 3: ROLES AND RESPONSIBILITIES
# ============================================================
page_break()
add_section_header("3. Roles and Responsibilities")
add_para(
    "The Department is responsible for providing legal guidance, documentation, and risk management oversight in "
    "support of the Company\u2019s business operations. The Department plays a central role in supporting the Company\u2019s "
    "leasing, operational, and governance activities while ensuring that legal risks are appropriately identified, "
    "evaluated, and managed."
)
add_para(
    "To promote clarity and operational efficiency, the Legal Department\u2019s responsibilities are organized into the "
    "categories set forth below. In certain areas, the Department is afforded sole discretion for a given business "
    "activity. In other areas, the Department merely provides guidance and support but does not serve as the operational "
    "\u201Cowner\u201D of the underlying business activity."
)

add_sub_header("Matters Owned and Managed by Legal Department")

add_para("**Lease Documentation and Legal Drafting**")
add_bullets([
    "Drafting and reviewing lease agreements, amendments, and related documentation",
    "Maintaining and updating standard lease forms and provisions",
    "Ensuring that lease documents reflect approved business terms and protect the Company\u2019s legal interests",
    "Managing internal document review and version control processes",
])

add_para("**Litigation Oversight**")
add_bullets([
    "Managing litigation involving the Company",
    "Coordinating with and supervising outside counsel",
    "Establishing litigation strategy and settlement posture in consultation with executive leadership",
    "Monitoring litigation budgets and controlling legal spending",
])

add_para("**Outside Counsel Management**")
add_bullets([
    "Selecting and maintaining a panel of approved outside counsel",
    "Issuing engagement instructions and defining scope of work",
    "Monitoring performance, staffing, and billing practices",
    "Ensuring adherence to the Company\u2019s billing guidelines and reporting expectations",
])

add_para("**Legal Risk Management and Reporting**")
add_bullets([
    "Maintaining the Company\u2019s Legal Risk Tracker",
    "Identifying and evaluating legal risks across the portfolio",
    "Reporting material legal matters to executive leadership",
    "Providing recommendations regarding mitigation strategies",
])

add_para("**Corporate Governance and Entity Management**")
add_bullets([
    "Maintaining records for corporate entities",
    "Supporting governance processes for the Company and its subsidiaries",
    "Coordinating entity formation, dissolution, and structural changes",
    "Maintaining organizational charts and ownership records",
])

add_para("**Policy Development and Legal Compliance**")
add_bullets([
    "Developing and maintaining internal legal policies and procedures",
    "Monitoring legal and regulatory developments relevant to the Company",
    "Advising internal teams regarding compliance obligations",
])

add_sub_header("Matters Where the Legal Department Provides Advisory Support or Collaboration")

add_para("**Lease/Document Negotiation for Asset Management Department**")
add_bullets([
    "Advising on legal implications of proposed business terms",
    "Recommending risk mitigation strategies",
    "Reviewing tenant revisions and proposed lease language",
])

add_para("**Tenant Defaults and Dispute Resolution for Executive Department**")
add_bullets([
    "Default notice procedures",
    "Enforcement options",
    "Settlement considerations",
    "Litigation strategies if disputes escalate",
])

add_para("**Insurance Claims for Asset Management Department**")
add_bullets([
    "Complex insurance coverage issues",
    "Insurer disputes",
    "Litigation arising from claims",
])

add_para("**Transactional Matters for Executive, Finance and Acquisitions Departments**")
add_bullets([
    "Acquisitions and dispositions",
    "Financing",
    "Joint ventures",
    "Material contracts",
])

add_sub_header("Matters Outside the Operational Responsibility of the Legal Department")
add_bullets([
    "Relationship management concerning tenants, vendors, contractors and third-party property management",
    "Rent collection and accounts receivable",
    "Property operations and maintenance",
    "Leasing strategy and market positioning",
    "Asset management decisions regarding tenant concessions or pricing",
])

# ============================================================
# SECTION 4: INTAKE & PRIORITIZATION
# ============================================================
page_break()
add_section_header("4. Intake & Prioritization Procedures")
add_para(
    "The Department receives requests from multiple internal stakeholders, including Asset Management, Property "
    "Management, Finance, and Executive Leadership. To ensure that legal matters are handled efficiently and "
    "consistently, all requests for legal services must follow the standardized intake and prioritization procedures "
    "outlined in this section."
)
add_para("The objectives of these procedures are to:")
add_numbered([
    "Ensure that legal requests are captured and tracked through a centralized system",
    "Establish clear and objective criteria for prioritizing work",
    "Provide visibility to leadership regarding legal workflow and capacity",
    "Ensure that matters affecting revenue, deadlines, or legal rights receive timely attention",
])

add_sub_header("Legal Request Intake")
add_para(
    "All requests for legal services related to leasing activities must be submitted through the Company\u2019s lease "
    "lifecycle management platform (currently Salesforce)."
)
add_para("The system automatically notifies the Legal Department when Asset Management submits a request for:")
add_bullets([
    "Drafting a new lease or lease amendment",
    "Reviewing a tenant-proposed lease draft or lease amendment",
    "Drafting notices relating to tenant defaults",
    "Reviewing other property-related legal documents",
])
add_para(
    "Using a centralized intake platform ensures that legal requests are documented, time-stamped, assigned an owner, "
    "and tracked through completion."
)
add_para("Asset Management or other requesting departments will be prompted for the following information when submitting a request:")
add_bullets([
    "Property name and address",
    "Tenant name",
    "Type of document requested",
    "First/Tenant-revised draft",
    "Lease/Renewal commencement date (if applicable)",
    "Lease Rent over the Lease Term",
])
add_para("Providing this information enables the Department to properly prioritize and allocate resources.")

add_sub_header("Guidance on Prioritization")
add_para(
    "The head of the Asset Management Department, and other authorized persons, will be given ultimate discretion "
    "for the priority of lease requests. However, the Company should aim to provide solutions that can automate these "
    "decisions based on a formula factoring in the information provided above. Using this information, the following "
    "business and operational considerations should be included in the priority analysis:"
)
add_bullets([
    "Requests affecting leases scheduled to commence in the near term",
    "Matters involving significant rent or strategic tenants",
    "Requests involving tenant revisions or active negotiations to maintain deal momentum",
    "Matters involving properties that are difficult to lease or strategically important to the portfolio",
    "Matters involving contractual deadlines, regulatory requirements, or litigation deadlines",
])

add_para("**High priority matters:**")
add_bullets([
    "New leases tied to upcoming rent commencement",
    "Tenant revisions requiring timely response",
    "Matters involving significant revenue impact",
    "Issues involving legal deadlines or risk of rights waiver",
    "Matters escalated by Executive Leadership",
])
add_para("These matters should generally receive attention within one to two business days, subject to complexity.")

add_para("**Medium priority matters:**")
add_bullets([
    "Lease amendments that do not affect immediate rent commencement",
    "Ongoing lease negotiations without immediate deadlines",
    "Routine document review requests (e.g., landlord lien waivers, NDAs, collateral access agreements, SNDAs, estoppel certificates, etc.)",
])
add_para("These matters are addressed in the normal legal workflow queue.")

add_para("**Low priority matters:**")
add_bullets([
    "Housekeeping amendments",
    "Internal document reviews",
    "Informational requests",
    "Long-term projects without time sensitivity",
])
add_para("These matters are handled as capacity permits.")

add_sub_header("Legal Workflow Dashboard")
add_para(
    "Salesforce will be programmed to compile legal requests into a centralized dashboard that provides visibility "
    "into pending legal requests, priority designations, responsible legal personnel, draft status and expected "
    "completion timelines."
)
add_para(
    "This dashboard helps ensure transparency across the leasing pipeline and allows leadership to monitor the flow "
    "of legal work across the organization."
)

add_sub_header("Monitoring and Updates")
add_para(
    "The Department, together with the Asset Management Department, should periodically review the status of open "
    "matters to ensure that priorities remain appropriate as circumstances evolve."
)
add_para(
    "While priority assignments are generally based on objective criteria noted above, certain matters may require "
    "adjustment based on business judgment. As previously noted, certain authorized persons may override the standard "
    "priority designation when necessary to reflect the Company\u2019s broader business needs and legal risk considerations."
)

# ============================================================
# SECTION 5: LEASING ASSET MANAGEMENT PROCEDURES
# ============================================================
page_break()
add_section_header("5. Leasing Asset Management Procedures")
add_para(
    "Leasing activity represents a core component of the Company\u2019s operations and revenue generation. The Legal "
    "Department supports Asset Management in connection with lease negotiations and documentation while ensuring that "
    "lease agreements appropriately protect the Company\u2019s legal and financial interests."
)
add_para("The procedures described in this section establish standardized workflows for:")
add_bullets([
    "Lease drafting and review",
    "Document collaboration between Legal and Asset Management",
    "Document version control",
    "Internal and external circulation of lease drafts",
])
add_para(
    "These procedures are designed to maintain document integrity, efficient collaboration, and clear accountability "
    "throughout the leasing process."
)

add_sub_header("Leasing Workflow Overview")
add_para("The leasing workflow generally involves the following steps:")
add_numbered([
    "Asset Management designates a SharePoint workspace to store draft Word documents (e.g., lease analysis, LOIs, etc.) and to allow internal review by the deal team.",
    "Asset Management submits a request for legal drafting or review through Salesforce\u2019s Leasing Process system.",
    "Legal prepares the initial draft or reviews the proposed form. If the request is to review a tenant-proposed or tenant-revised form, Asset Management should first provide internal comments and revisions in the document using the \u201CComment\u201D and \u201CTracked Changes\u201D functions. Legal will (i) revise and provide comments to Asset Management in the Word document using the \u201C@mention\u201D function and (ii) review and incorporate internal comments and revisions, as applicable. Drafting-related communications should be made using these functions, in lieu of e-mail.",
    "Once internal review is completed, Legal or Asset Management will (i) generate a comparison using the \u201CCompare\u201D function in Word of the draft against the tenant-proposed form, if applicable, and (ii) download a local copy of the SharePoint draft and circulate the draft to tenant counsel or the tenant\u2019s representative, as applicable. Asset Management will then indicate, through Salesforce\u2019s Leasing Process system, that the lease or lease amendment has proceeded to negotiation phase.",
    "Lease drafts are revised as negotiations progress. Every initial and subsequent external draft will be saved as a separate Word document in accordance with applicable naming convention (See Naming Convention below).",
    "Final documents are prepared for signature and distribution by the applicable Leasing Administrator from Asset Management and stored in ShareFile (not to be confused with SharePoint).",
])

add_sub_header("Role of Asset Management")
add_para(
    "Asset Management is responsible for, and has final authority over, the commercial negotiation of lease terms "
    "and for maintaining the tenant (and broker, if applicable) relationship throughout the negotiation process."
)
add_para("Asset Management responsibilities include:")
add_bullets([
    "Negotiating economic terms with tenants",
    "Coordinating with tenants and brokers regarding business terms",
    "Routing drafts between the Legal Department, tenants and brokers",
    "Confirming that negotiated business terms are accurately reflected in the lease documentation",
])
add_para(
    "On some occasions, Asset Management may be permitted to draft or document a business term in a legal document; "
    "however, the Department will have final approval of the document. If for any reason, typically due to time-sensitivity, "
    "the Department does not have the opportunity to review a document, Asset Management may be permitted to transmit "
    "provided, that, in its transmission correspondence, Asset Management expressly reserves the right for final review "
    "by the Company\u2019s Legal Department."
)
add_para("The Department supports Asset Management by ensuring that negotiated terms are properly documented and that legal risks are appropriately addressed.")
add_para("Legal responsibilities include:")
add_bullets([
    "Drafting lease agreements and amendments",
    "Reviewing tenant-proposed lease forms",
    "Ensuring consistency with the Company\u2019s standard lease provisions",
    "Advising Asset Management regarding legal implications of revisions made by tenant and/or Asset Management",
    "Maintaining document version control",
    "Coordinating the internal review process prior to external circulation",
])

add_sub_header("SharePoint")
add_para("All lease-related documents must be stored in the designated SharePoint workspace for the applicable transaction.")
add_para("The SharePoint workspace enables:")
add_bullets([
    "Centralized document storage",
    "Version history tracking for auditing drafting decisions",
    "Internal collaboration",
    "Preservation of document integrity",
])
add_para("Maintaining a centralized repository ensures that all members of the deal team are working from the most current version of the document.")

add_sub_header("Document Naming Convention")
add_para("All draft lease documents must follow the Company\u2019s standardized naming convention to ensure consistency and clarity.")
add_para("The standard naming format is:")
add_para("**[Document Type] \u2013 [Tenant Name] \u2013 [Property Address] \u2013 PLYM Draft**")
add_para(
    "Subsequent versions, whether or not proposed by the Company, should be numbered sequentially. Tenant-proposed forms "
    "will constitute the first version of any draft to be reviewed internally by the deal team. The most recent version of "
    "the document, as reflected in the version sequence, shall be deemed the operative draft for external circulation to the "
    "tenant or, if finalized, for execution."
)
add_para("Lease \u2013 ABC Logistics \u2013 100 Industrial Drive \u2013 PLYM Draft v2", italic=True)
add_para("Lease \u2013 ABC Logistics \u2013 100 Industrial Drive \u2013 PLYM Draft v3", italic=True)
add_para("Consistent naming conventions help prevent confusion and allow the deal team to easily identify the most recent version of a document.")

add_sub_header("Tracked Changes Protocol")
add_para(
    "Tracked changes must remain enabled during the lease negotiation process to preserve visibility into both internal "
    "and external revisions and maintain a clear record of document changes. Tracked changes should not be removed before "
    "internal review is completed. The Department should ensure that internal comments are resolved prior to approval of "
    "external circulation. After deleting resolved comments, the Department will generate a comparison of the initial "
    "version of a tenant-proposed form, as applicable, and the most recent version revised by the Department and the deal team."
)
add_para("Maintaining tracked changes ensures transparency and facilitates efficient negotiation.")

add_sub_header("Continuous Improvement of Lease Documentation")
add_para(
    "The Legal Department periodically reviews leasing activity to identify opportunities to improve the Company\u2019s "
    "standard lease forms and negotiation practices."
)
add_para("Areas of review may include:")
add_bullets([
    "Frequently negotiated provisions",
    "Recurring tenant concerns",
    "Legal developments affecting lease provisions",
    "Operational feedback from Asset Management",
])
add_para("These reviews help ensure that the Company\u2019s leasing documentation evolves to reflect market conditions, operational needs, and risk management priorities.")

# ============================================================
# SECTION 6: LITIGATION & OUTSIDE COUNSEL MANAGEMENT
# ============================================================
page_break()
add_section_header("6. Litigation & Outside Counsel Management")
add_para(
    "The Legal Department is responsible for overseeing all litigation and dispute-related matters involving the Company. "
    "This includes managing legal risk, coordinating with outside counsel, controlling legal expenses, and ensuring that "
    "litigation strategy aligns with the Company\u2019s business objectives."
)
add_para("The procedures set forth in this section are designed to:")
add_bullets([
    "Ensure consistent and centralized management of litigation",
    "Maintain visibility into legal exposure across the portfolio",
    "Promote cost-effective use of outside counsel",
    "Support timely and informed decision-making by executive leadership",
])

add_sub_header("Litigation Intake and Tracking")
add_para(
    "All litigation and dispute matters must be logged in the Legal Risk Tracker immediately upon receipt of a demand, "
    "complaint, or notice of claim. The Department will assign each matter a unique identifier and designate a responsible "
    "attorney or legal professional."
)
add_para("For each matter, the following information must be documented:")
add_bullets([
    "Parties involved",
    "Nature of the claim or dispute",
    "Property or entity affected",
    "Date received and key deadlines",
    "Assigned outside counsel (if applicable)",
    "Estimated financial exposure",
    "Litigation strategy and settlement posture",
])

add_sub_header("Outside Counsel Selection and Engagement")
add_para(
    "The Department maintains a panel of approved outside counsel firms, selected based on relevant expertise, "
    "jurisdictional coverage, responsiveness, and cost efficiency. All outside counsel engagements must be initiated "
    "by the Department."
)
add_para("Prior to engagement, the Department will:")
add_numbered([
    "Define the scope of work and expected deliverables",
    "Confirm applicable billing rates and rate caps (see Attachment C: Outside Counsel Billing Guidelines)",
    "Establish a matter budget and obtain required approvals for engagements exceeding established thresholds (see Attachment A: Delegation of Authority Matrix)",
    "Issue a written engagement letter specifying terms, staffing expectations, and reporting requirements",
])

add_sub_header("Litigation Budget and Spend Management")
add_para(
    "Outside counsel must submit matter budgets at the outset of each engagement. The Department monitors actual "
    "spend against budgets on a monthly basis."
)
add_para("Budget oversight procedures include:")
add_bullets([
    "Monthly invoice review and approval by the responsible Department attorney",
    "Comparison of actual spend to approved budget with variance analysis",
    "Quarterly reporting of aggregate legal spend to the CFO",
    "Pre-approval required for expenditures exceeding the original matter budget by more than 15%",
])

add_sub_header("Litigation Strategy and Settlement Authority")
add_para(
    "The Department develops litigation strategy in consultation with executive leadership. Settlement authority is "
    "governed by the Delegation of Authority Matrix (Attachment A). All settlement proposals must be documented in "
    "writing and include an analysis of the financial and strategic implications of the proposed resolution."
)

add_sub_header("Litigation Reporting")
add_para("The Department provides periodic litigation updates to executive leadership, including:")
add_bullets([
    "Status of all active matters",
    "Changes in estimated exposure",
    "Key developments and upcoming deadlines",
    "Budget vs. actual spend analysis",
    "Recommended actions or strategic decisions required",
])
add_para("Litigation reporting is integrated with the Legal Risk Tracker and the Legal Risk Dashboard (see Section 10 and Attachment B).")

add_sub_header("Post-Litigation Review")
add_para(
    "Upon resolution of significant matters, the Department conducts a post-litigation review to identify lessons learned "
    "and opportunities for process improvement. This review may include evaluation of outside counsel performance, "
    "assessment of litigation outcomes relative to initial exposure estimates, and identification of recurring issues "
    "that may warrant changes to lease forms, operational procedures, or risk management practices."
)

# ============================================================
# SECTION 7: CORPORATE GOVERNANCE & ENTITY MANAGEMENT
# ============================================================
page_break()
add_section_header("7. Corporate Governance & Entity Management")
add_para(
    "The Legal Department is responsible for overseeing the Company\u2019s corporate governance framework and maintaining "
    "accurate and complete records for all legal entities within the Company\u2019s organizational structure."
)
add_para("These procedures are intended to:")
add_bullets([
    "Ensure compliance with applicable corporate and REIT requirements",
    "Maintain accurate records of entity structure and authority",
    "Support financing, transactional, and operational needs",
    "Preserve institutional knowledge of the Company\u2019s legal structure",
])

add_sub_header("A. Entity Management")
add_para("The Legal Department maintains a centralized record of all Company entities in SharePoint, including:")
add_bullets([
    "Legal entity names and jurisdictions of formation",
    "Ownership structure and organizational charts",
    "Governing documents (e.g., operating agreements, bylaws)",
    "Registered agents and principal offices",
    "Qualification and good standing status",
])
add_para("Legal is responsible for coordinating:")
add_bullets([
    "Entity formation and dissolution",
    "Amendments to governing documents",
    "Intercompany structuring changes",
])
add_para(
    "Entity records must be reviewed and updated at least annually, or upon the occurrence of a material event "
    "(e.g., acquisition, disposition, financing, or restructuring). The Department maintains a master entity "
    "spreadsheet that serves as the primary reference for all entity-related inquiries."
)

add_sub_header("B. Governance Documentation")
add_para("The Legal Department maintains and organizes governance records, including:")
add_bullets([
    "Member and board approvals",
    "Written consents and resolutions",
    "Officer and manager appointments",
    "Delegation of authority documentation",
])
add_para(
    "All governance actions must be properly documented and retained in the Company\u2019s document management system "
    "(SharePoint). The Department maintains a governance calendar to track recurring approval requirements, annual "
    "meeting obligations, and filing deadlines."
)

add_sub_header("C. Authority and Approvals")
add_para(
    "The Legal Department supports the implementation and maintenance of the Company\u2019s delegation of authority "
    "framework (see Attachment A: Delegation of Authority Matrix)."
)
add_para("Responsibilities include:")
add_bullets([
    "Confirming authority for execution of contracts and leases",
    "Coordinating approvals for material transactions",
    "Maintaining records of authorized signatories",
    "Verifying that approval thresholds are followed prior to document execution",
])
add_para(
    "The Department will circulate an updated list of authorized signatories to relevant departments no less than "
    "annually and upon any change in authorized personnel."
)

add_sub_header("D. Compliance and Filings")
add_para("The Legal Department coordinates with internal and external service providers to ensure:")
add_bullets([
    "Timely filing of annual reports and required state filings",
    "Maintenance of good standing for all entities",
    "Compliance with REIT-related legal requirements",
])
add_para(
    "The Department utilizes an external registered agent service provider for state filings and good standing "
    "maintenance. The Department maintains a compliance calendar with all recurring filing deadlines, organized by "
    "entity and jurisdiction. Filing status is reviewed quarterly and reported to the CFO."
)

# ============================================================
# SECTION 8: CONFIDENTIALITY & INFORMATION HANDLING
# ============================================================
page_break()
add_section_header("8. Confidentiality & Information Handling")
add_para(
    "The Legal Department handles sensitive and confidential information, including privileged communications, "
    "litigation materials, and proprietary business information. Proper handling of such information is critical "
    "to protecting the Company\u2019s legal position and maintaining confidentiality obligations."
)

add_sub_header("A. Confidential Information")
add_para("Confidential information includes, but is not limited to:")
add_bullets([
    "Legal advice and attorney-client communications",
    "Litigation strategy and work product",
    "Tenant disputes and negotiations",
    "Regulatory matters",
    "Transaction-related information",
])

add_sub_header("B. Attorney-Client Privilege")
add_para(
    "Communications involving the Legal Department that are intended to provide or request legal advice may be "
    "protected by attorney-client privilege."
)
add_para("Employees should:")
add_bullets([
    "Limit distribution of privileged communications",
    "Avoid forwarding privileged emails outside the Company without Legal approval",
    "Clearly label privileged communications where appropriate",
])
add_para(
    "Privileged communications should be labeled with the following header: \u201CPRIVILEGED AND CONFIDENTIAL \u2013 "
    "ATTORNEY-CLIENT COMMUNICATION.\u201D The Department will periodically remind internal stakeholders of privilege "
    "preservation obligations through training and written guidance."
)

add_sub_header("C. Document Handling and Retention")
add_para(
    "Legal documents should be stored in approved systems (e.g., SharePoint) and should not be stored locally or "
    "outside authorized platforms. Access to sensitive materials should be limited to individuals with a business need."
)
add_para("The Department maintains a document retention schedule, organized by category:")

add_table(
    ["Document Category", "Retention Period", "Storage Location"],
    [
        ["Executed Leases", "Life of lease + 7 years", "ShareFile"],
        ["Lease Drafts and Negotiations", "Life of lease + 3 years", "SharePoint"],
        ["Litigation Files", "Resolution + 7 years", "SharePoint / Outside Counsel"],
        ["Corporate Governance Records", "Permanent", "SharePoint"],
        ["Outside Counsel Invoices", "7 years", "SharePoint / Finance"],
        ["General Correspondence", "3 years", "SharePoint"],
        ["Regulatory Filings", "Permanent", "SharePoint"],
    ],
    [2.2, 2.2, 2.1]
)

add_sub_header("D. Litigation Hold Procedures")
add_para(
    "Upon identification of a matter that may result in litigation or regulatory investigation, the Department "
    "will implement a litigation hold to preserve all potentially relevant documents and communications."
)
add_para("Litigation hold procedures include:")
add_numbered([
    "The Department issues a written litigation hold notice to all relevant custodians, identifying the matter and the categories of documents to be preserved.",
    "Custodians must acknowledge receipt of the hold notice in writing.",
    "The Department coordinates with IT to suspend any automated deletion or archiving processes for documents within the scope of the hold.",
    "The hold remains in effect until the Department issues a written release notice.",
    "The Department maintains a log of all active and released litigation holds.",
])

add_sub_header("E. External Communications")
add_para("Only authorized personnel may communicate with:")
add_bullets([
    "Outside counsel",
    "Opposing counsel",
    "Regulatory authorities",
])
add_para("All external legal communications should be coordinated through the Legal Department.")

# ============================================================
# SECTION 9: COMMUNICATION & ESCALATION PROTOCOLS
# ============================================================
page_break()
add_section_header("9. Communication & Escalation Protocols")
add_para(
    "Clear communication and timely escalation are essential to effective legal risk management. This section defines "
    "when and how matters should be escalated to the Legal Department and executive leadership."
)

add_sub_header("A. Routine Communication")
add_para("The Legal Department works collaboratively with:")
add_bullets([
    "Asset Management (leasing and tenant matters)",
    "Property Management (operational issues and incidents)",
    "Finance (reserves, transactions, governance)",
])
add_para("Standing interdepartmental meetings should be held at the following cadence:")

add_table(
    ["Meeting", "Participants", "Frequency", "Purpose"],
    [
        ["Leasing Pipeline Review", "Legal + Asset Mgmt", "Weekly", "Review pending lease requests, priority alignment"],
        ["Litigation/Risk Update", "Legal + Executive Leadership", "Bi-weekly", "Active matters, exposure, strategic decisions"],
        ["Finance Coordination", "Legal + Finance/CFO", "Monthly", "Reserves, legal spend, governance filings"],
        ["Property Ops Sync", "Legal + Property Mgmt", "As needed", "Incidents, insurance claims, vendor disputes"],
    ],
    [1.5, 1.5, 1.2, 2.3]
)

add_sub_header("B. Escalation Criteria")
add_para("Matters must be escalated to the Legal Department when they involve:")
add_bullets([
    "Potential or actual litigation",
    "Tenant bankruptcy or insolvency",
    "Material financial exposure",
    "Regulatory or governmental inquiries",
    "Significant personal injury or property damage",
    "Disputes that may impact Company rights or obligations",
])

add_sub_header("C. Escalation Timelines")
add_para("The following escalation timelines apply:")

add_table(
    ["Event Type", "Escalation Timeline", "Escalate To"],
    [
        ["Receipt of lawsuit or demand letter", "Immediately (same business day)", "General Counsel"],
        ["Regulatory inquiry or subpoena", "Immediately (same business day)", "General Counsel + CEO"],
        ["Tenant bankruptcy filing", "Within 24 hours", "General Counsel + Asset Mgmt Lead"],
        ["Significant personal injury on property", "Immediately", "General Counsel + Property Mgmt Lead"],
        ["Material financial exposure (> $250,000)", "Within 24 hours", "General Counsel + CFO"],
        ["Potential privilege waiver or breach", "Immediately", "General Counsel"],
        ["Insurance coverage dispute", "Within 48 hours", "General Counsel + Asset Mgmt Lead"],
    ],
    [2.2, 2.2, 2.1]
)

add_sub_header("D. Executive Escalation")
add_para("The Legal Department escalates matters to executive leadership when they involve:")
add_bullets([
    "Significant financial exposure",
    "Reputational risk",
    "Portfolio-wide implications",
    "Strategic business impact",
])
add_para(
    "Executive escalation should include a written summary of the matter, current status, recommended course of action, "
    "and estimated financial impact. The Department will prepare escalation memoranda using a standardized format to "
    "ensure consistency."
)

add_sub_header("E. Legal Updates and Reporting")
add_para("The Legal Department provides periodic updates to executive leadership, including:")
add_bullets([
    "Status of significant matters",
    "Litigation exposure",
    "Key risks and developments",
])
add_para("The Legal Risk Tracker and Legal Risk Dashboard (see Section 10 and Attachment B) are used as primary tools for these updates.")

# ============================================================
# SECTION 10: LEGAL RISK TRACKING & REPORTING
# ============================================================
page_break()
add_section_header("10. Legal Risk Tracking & Reporting")
add_para(
    "The Legal Department maintains a Legal Risk Tracker to provide visibility into legal matters affecting the "
    "Company\u2019s portfolio."
)
add_para("The tracker serves as a centralized tool for:")
add_bullets([
    "Monitoring litigation and disputes",
    "Tracking tenant defaults and risk exposure",
    "Identifying trends and recurring issues",
    "Supporting executive decision-making",
])

add_sub_header("A. Risk Categories")
add_para("Legal matters are categorized into:")
add_bullets([
    "Litigation",
    "Tenant defaults",
    "Bankruptcy",
    "Insurance claims",
    "Regulatory matters",
    "Environmental issues",
])

add_sub_header("B. Risk Assessment Methodology")
add_para("Each matter is evaluated based on a probability-impact framework:")

add_table(
    ["Rating", "Probability", "Financial Impact"],
    [
        ["High", "Likely (>70%)", "> $500,000"],
        ["Medium", "Possible (30\u201370%)", "$100,000 \u2013 $500,000"],
        ["Low", "Unlikely (<30%)", "< $100,000"],
    ],
    [2.2, 2.2, 2.1]
)

add_para("Each matter also receives a qualitative assessment considering:")
add_bullets([
    "Strategic or precedent risk",
    "Potential reputational impact",
    "Insurance coverage availability",
    "Strength of the Company\u2019s legal position",
])
add_para("Risk ratings are reviewed and updated at each reporting cycle. Changes in risk rating must be documented with supporting rationale.")

add_sub_header("C. Reporting Cadence")
add_para("The Legal Risk Tracker is updated on a regular basis:")
add_bullets([
    "High priority matters: weekly",
    "Medium priority matters: monthly",
    "Low priority matters: quarterly",
])
add_para(
    "Updated reports are shared with executive leadership in advance of scheduled legal update meetings. The Legal "
    "Risk Dashboard (Attachment B) is prepared monthly for the CFO and quarterly for board-level reporting."
)

add_sub_header("D. Integration with Reserve Analysis")
add_para(
    "Risk data from the Legal Risk Tracker is used to inform the Company\u2019s reserve analysis, prepared in coordination "
    "with Finance. The Department provides estimated exposure ranges for all active matters rated Medium or High, updated "
    "quarterly or upon any material change in circumstances."
)

add_sub_header("E. Use of Risk Data")
add_para("Risk data is used to:")
add_bullets([
    "Inform reserve decisions",
    "Support strategic planning",
    "Identify areas for risk mitigation",
    "Improve lease documentation and processes",
])

# ============================================================
# SECTION 11: TRAINING & CONTINUOUS IMPROVEMENT
# ============================================================
page_break()
add_section_header("11. Training & Continuous Improvement")
add_para(
    "The Legal Department supports ongoing training and development to improve collaboration, reduce risk, and "
    "enhance operational efficiency."
)

add_sub_header("A. Internal Training Program")
add_para("The Legal Department conducts training for internal stakeholders on the following schedule:")

add_table(
    ["Audience", "Frequency", "Topics"],
    [
        ["Asset Management", "Quarterly", "Lease provisions, negotiation practices, document workflows"],
        ["Property Management", "Quarterly", "Incident reporting, insurance claims, escalation procedures"],
        ["Finance", "Semi-annually", "Legal risk reporting, reserve integration, governance processes"],
        ["All Departments", "Annually", "Confidentiality obligations, privilege preservation, policy updates"],
        ["New Hires (Legal)", "Upon onboarding", "Systems training, standard forms, naming conventions, workflows"],
    ],
    [1.8, 1.5, 3.2]
)

add_sub_header("B. New Hire Onboarding")
add_para("New Department personnel complete the following onboarding checklist:")
add_numbered([
    "Review of this Manual and all attachments",
    "Overview of the Company\u2019s entity structure and portfolio",
    "Salesforce system training (intake, workflow, dashboard)",
    "SharePoint training (document storage, collaboration, naming conventions)",
    "Introduction to standard lease forms and key provisions",
    "Review of active litigation and risk tracker",
    "Introduction to outside counsel panel and billing guidelines",
    "Review of confidentiality and privilege protocols",
])

add_sub_header("C. Training Materials")
add_para(
    "Training materials, including presentation decks, reference guides, and recorded sessions, are stored in a "
    "dedicated SharePoint folder maintained by the Department. Materials are updated as procedures, systems, or "
    "standard forms change."
)

add_sub_header("D. Legal Updates")
add_para(
    "The Legal Department monitors legal and market developments and provides updates to internal stakeholders as "
    "appropriate. Updates may be distributed via email memoranda, incorporated into training sessions, or presented "
    "at standing interdepartmental meetings."
)

add_sub_header("E. Continuous Improvement")
add_para("The Legal Department regularly reviews:")
add_bullets(["Leasing processes", "Litigation outcomes", "Recurring issues"])
add_para("To identify opportunities to:")
add_bullets(["Improve standard lease forms", "Streamline workflows", "Reduce legal exposure"])
add_para(
    "Post-litigation and post-transaction reviews are conducted for significant matters to identify lessons learned. "
    "Findings are documented and, where applicable, incorporated into updated procedures, training materials, or standard forms."
)

# ============================================================
# SECTION 12: TECHNOLOGY & WORKFLOW OPTIMIZATION
# ============================================================
page_break()
add_section_header("12. Technology & Workflow Optimization")
add_para(
    "The Legal Department utilizes technology to improve efficiency, enhance visibility, and support scalable operations."
)

add_sub_header("A. Core Systems")
add_para("The following systems are currently used by the Department:")

add_table(
    ["System", "Primary Function", "Owner", "Users"],
    [
        ["Salesforce", "Lease intake, workflow tracking, pipeline dashboard", "Legal / Asset Mgmt", "Legal, Asset Mgmt, Executive"],
        ["SharePoint", "Document collaboration, version control, storage", "Legal / IT", "Legal, Asset Mgmt, Finance"],
        ["ShareFile", "Executed document storage and external sharing", "Asset Mgmt", "Legal, Asset Mgmt"],
        ["Legal Risk Tracker", "Litigation and risk monitoring, reporting", "Legal", "Legal, Executive, CFO"],
    ],
    [1.3, 2.5, 1.3, 1.4]
)

add_para("System access is provisioned based on role and business need. The Department coordinates with IT to ensure appropriate access controls are maintained.")

add_sub_header("B. System Integration")
add_para("The Legal Department works to ensure that:")
add_bullets([
    "Legal workflows are integrated with operational systems",
    "Data is captured consistently across platforms",
    "Information is accessible to relevant stakeholders",
    "Duplicate data entry is minimized through system integration where feasible",
])
add_para(
    "The Department collaborates with IT and Finance to ensure that legal data (e.g., lease milestones, litigation "
    "reserves, entity records) is accurately reflected in the Company\u2019s financial and operational reporting systems."
)

add_sub_header("C. Evaluation of New Tools")
add_para("The Legal Department may evaluate additional tools to improve:")
add_bullets([
    "Lease data extraction and analysis",
    "Document automation",
    "Risk tracking and reporting",
    "Workflow automation",
])
add_para("New technology evaluations should consider the following criteria:")
add_bullets([
    "Integration capability with existing systems (Salesforce, SharePoint)",
    "Data security and access control standards",
    "Cost-benefit analysis and return on investment",
    "Ease of adoption and training requirements",
    "Vendor reliability and support",
])
add_para(
    "The Department will coordinate with IT and Finance before procuring or implementing new technology solutions. "
    "Material technology investments require approval in accordance with the Company\u2019s standard procurement process."
)

add_sub_header("D. Data Integrity")
add_para("Maintaining accurate and consistent data across systems is critical to effective legal operations and reporting. The Department is responsible for:")
add_bullets([
    "Periodic audits of data accuracy in Salesforce and SharePoint",
    "Ensuring that lease milestone dates, entity records, and litigation data are current",
    "Reporting data discrepancies to the relevant system owner for correction",
    "Maintaining consistent naming conventions and folder structures across all platforms",
])

add_sub_header("E. Change Management")
add_para(
    "System upgrades, migrations, or workflow changes that affect the Department\u2019s operations must be coordinated "
    "with the Department in advance. The Department will participate in testing and training prior to implementation "
    "of material system changes to ensure continuity of legal operations."
)

# ============================================================
# SECTION 13: POLICY REVIEW AND UPDATES
# ============================================================
page_break()
add_section_header("13. Policy Review and Updates")
add_para("This Manual is intended to serve as a living document that evolves with the Company\u2019s operations and legal requirements.")

add_sub_header("A. Periodic Review")
add_para("The Legal Department will review this Manual annually, with the review cycle aligned to the Company\u2019s fiscal year. Reviews should ensure alignment with:")
add_bullets([
    "Changes in Company structure or strategy",
    "System implementations or upgrades",
    "Regulatory developments",
    "Operational needs",
])
add_para("The annual review should be completed within the first quarter of each fiscal year.")

add_sub_header("B. Updates and Revisions")
add_para(
    "The Legal Department may update this Manual as necessary. Material changes must be approved by the General Counsel "
    "and communicated to relevant stakeholders."
)
add_para("The following version control practices apply:")
add_bullets([
    "Each revision is assigned a version number and effective date",
    "A change log is maintained documenting all material revisions (see below)",
    "The current version of the Manual is stored in SharePoint in a designated location accessible to Department personnel",
    "Superseded versions are archived but retained for reference",
])

add_sub_header("C. Change Log")

add_table(
    ["Version", "Date", "Description of Change", "Approved By"],
    [
        ["1.0", "[Date]", "Initial publication", "[Name/Title]"],
        ["", "", "", ""],
        ["", "", "", ""],
    ],
    [0.9, 1.1, 3.3, 1.2]
)

add_sub_header("D. Implementation")
add_para(
    "The Legal Department is responsible for implementing and maintaining compliance with the procedures outlined in "
    "this Manual. All Department personnel are expected to review and adhere to the current version of this Manual."
)

# ============================================================
# ATTACHMENT A: DELEGATION OF AUTHORITY MATRIX
# ============================================================
page_break()
add_section_header("Attachment A: Delegation of Authority Matrix")
add_para(
    "The following matrix establishes approval authority for legal and transactional matters. All dollar thresholds are "
    "subject to review and adjustment by executive leadership. Amounts exceeding the highest tier require Board approval."
)

add_table(
    ["Action / Document Type", "Tier 1: GC", "Tier 2: GC + CFO", "Tier 3: GC + CEO", "Tier 4: Board"],
    [
        ["New Lease Execution", "< $[X] annual rent", "$[X] \u2013 $[Y]", "$[Y] \u2013 $[Z]", "> $[Z]"],
        ["Lease Amendment", "< $[X] impact", "$[X] \u2013 $[Y]", "$[Y] \u2013 $[Z]", "> $[Z]"],
        ["Lease Termination / Surrender", "< $[X]", "$[X] \u2013 $[Y]", "$[Y] \u2013 $[Z]", "> $[Z]"],
        ["Litigation Settlement", "< $[X]", "$[X] \u2013 $[Y]", "$[Y] \u2013 $[Z]", "> $[Z]"],
        ["Outside Counsel Engagement", "< $[X] budget", "$[X] \u2013 $[Y]", "> $[Y]", "N/A"],
        ["Vendor / Service Contracts", "< $[X]", "$[X] \u2013 $[Y]", "$[Y] \u2013 $[Z]", "> $[Z]"],
        ["Entity Formation / Dissolution", "GC approval", "GC + CFO", "GC + CEO", "Board"],
        ["Acquisition / Disposition", "N/A", "N/A", "GC + CEO", "Board"],
        ["Financing / Loan Documents", "N/A", "GC + CFO", "GC + CEO", "Board"],
        ["Insurance Claim Settlement", "< $[X]", "$[X] \u2013 $[Y]", "> $[Y]", "N/A"],
    ],
    [1.8, 1.3, 1.3, 1.3, 0.8]
)

add_para("**Notes:**")
add_bullets([
    "Dollar thresholds denoted as $[X], $[Y], and $[Z] are to be established by executive leadership and populated prior to implementation.",
    "All approvals must be documented in writing (email confirmation is acceptable for Tier 1 matters).",
    "Emergency authority: In the absence of the General Counsel, the CEO may exercise Tier 1 and Tier 2 authority for time-sensitive matters, with subsequent documentation.",
    "This matrix should be reviewed annually in conjunction with the Manual review (Section 13).",
])

# ============================================================
# ATTACHMENT B: LEGAL RISK DASHBOARD
# ============================================================
page_break()
add_section_header("Attachment B: Legal Risk Dashboard")
add_para(
    "The Legal Risk Dashboard is a summary reporting tool designed for executive leadership and board-level consumption. "
    "It is prepared monthly by the Legal Department and distributed in advance of scheduled legal update meetings."
)

add_para("**Section 1: Active Matters Summary**")
add_table(
    ["Category", "High Risk", "Medium Risk", "Low Risk", "Total"],
    [
        ["Litigation", "[#]", "[#]", "[#]", "[#]"],
        ["Tenant Defaults", "[#]", "[#]", "[#]", "[#]"],
        ["Bankruptcy", "[#]", "[#]", "[#]", "[#]"],
        ["Insurance Claims", "[#]", "[#]", "[#]", "[#]"],
        ["Regulatory", "[#]", "[#]", "[#]", "[#]"],
        ["Environmental", "[#]", "[#]", "[#]", "[#]"],
        ["TOTAL", "[#]", "[#]", "[#]", "[#]"],
    ],
    [1.8, 1.3, 1.3, 1.1, 1.0]
)

add_para("**Section 2: Financial Exposure Summary**")
add_table(
    ["Risk Level", "Number of Matters", "Estimated Exposure Range", "Current Reserves"],
    [
        ["High", "[#]", "$[X] \u2013 $[Y]", "$[amount]"],
        ["Medium", "[#]", "$[X] \u2013 $[Y]", "$[amount]"],
        ["Low", "[#]", "$[X] \u2013 $[Y]", "$[amount]"],
        ["TOTAL", "[#]", "$[X] \u2013 $[Y]", "$[amount]"],
    ],
    [1.6, 1.6, 1.7, 1.6]
)

add_para("**Section 3: Key Developments**")
add_para("This section contains a narrative summary of material developments during the reporting period, including:")
add_bullets([
    "New matters opened",
    "Matters resolved or settled",
    "Significant changes in risk rating or exposure",
    "Upcoming deadlines or decisions requiring executive action",
    "Trends or recurring issues warranting attention",
])

add_para("**Section 4: Legal Spend Summary**")
add_table(
    ["Category", "YTD Actual", "YTD Budget", "Variance", "Forecast"],
    [
        ["Litigation", "$[X]", "$[X]", "$[X]", "$[X]"],
        ["Transactional", "$[X]", "$[X]", "$[X]", "$[X]"],
        ["Regulatory/Compliance", "$[X]", "$[X]", "$[X]", "$[X]"],
        ["TOTAL", "$[X]", "$[X]", "$[X]", "$[X]"],
    ],
    [1.8, 1.3, 1.3, 1.1, 1.0]
)

add_para("**Section 5: Action Items**")
add_para("Decisions or approvals requested from executive leadership or the Board, with recommended actions and deadlines.")

# ============================================================
# ATTACHMENT C: OUTSIDE COUNSEL BILLING GUIDELINES
# ============================================================
page_break()
add_section_header("Attachment C: Outside Counsel Billing Guidelines")
add_para(
    "These guidelines apply to all outside counsel retained by PIR Industrial LLC (the \u201CCompany\u201D). Outside "
    "counsel are expected to review and adhere to these guidelines as a condition of engagement. The Legal Department "
    "reserves the right to modify these guidelines at any time."
)

add_sub_header("1. Billing Rates")
add_bullets([
    "Rate schedules must be agreed upon in writing prior to commencement of work.",
    "Annual rate increases, if any, must be approved in advance by the General Counsel.",
    "Maximum hourly rates by level: Partner: $[X]/hr; Senior Associate: $[X]/hr; Associate: $[X]/hr; Paralegal: $[X]/hr.",
    "Blended or alternative fee arrangements are encouraged for appropriate matters and must be approved by the General Counsel.",
])

add_sub_header("2. Staffing Requirements")
add_bullets([
    "Matters must be staffed efficiently. The Company expects lean staffing appropriate to the complexity of the matter.",
    "The responsible partner must be identified in the engagement letter and may not be changed without prior approval.",
    "Work should be delegated to the lowest-cost attorney capable of performing it competently.",
    "The Company will not pay for more than two attorneys attending the same deposition, hearing, or meeting without prior approval.",
    "Summer associates and first-year associates may not bill time to Company matters without prior approval.",
])

add_sub_header("3. Billing Format and Requirements")
add_bullets([
    "Invoices must be submitted monthly, within 45 days of the end of the billing period.",
    "Each invoice must include: matter name and number, description of services by timekeeper, hours worked, hourly rate, and total fees.",
    "Time entries must be recorded in increments of no greater than one-tenth (0.1) of an hour.",
    "Block billing (grouping multiple tasks into a single time entry) is prohibited.",
    "Vague descriptions (e.g., \u201Cresearch,\u201D \u201Creview documents,\u201D \u201Cattend to matter\u201D) are not acceptable. Each entry must describe the specific work performed.",
])

add_sub_header("4. Prohibited Charges")
add_para("The following charges will not be reimbursed:")
add_bullets([
    "Internal administrative and overhead costs (word processing, secretarial support, filing, copying for internal use)",
    "Database or legal research service access fees (e.g., Westlaw, LexisNexis subscriptions)",
    "Travel time (unless pre-approved in writing for out-of-jurisdiction matters)",
    "First-class or business-class airfare",
    "Alcohol",
    "Charges for internal firm meetings, training, or professional development",
    "Rush charges for printing, copying, or courier services not requested by the Company",
])

add_sub_header("5. Expense Pre-Approval")
add_para("The following expenses require written pre-approval from the Department:")
add_bullets([
    "Expert witnesses and consultants",
    "Travel and lodging",
    "Outsourced discovery or document review services",
    "Filing fees exceeding $1,000",
    "Any single expense exceeding $2,500",
])

add_sub_header("6. Matter Budgets")
add_bullets([
    "A matter budget is required at the outset of each engagement, broken down by phase of work.",
    "The budget must be approved by the Department before substantive work begins.",
    "If actual spend is projected to exceed the approved budget by more than 15%, outside counsel must notify the Department in writing and obtain approval for a revised budget before incurring additional charges.",
    "Quarterly budget-to-actual reports must be submitted with each invoice.",
])

add_sub_header("7. Reporting Requirements")
add_para("Outside counsel must provide:")
add_bullets([
    "Monthly status reports summarizing activity, developments, and next steps",
    "Prompt notification of any material development (adverse ruling, settlement demand, deadline change)",
    "Quarterly budget-to-actual analysis",
    "Early case assessment within 30 days of engagement, including estimated exposure range, recommended strategy, and projected timeline",
])

add_sub_header("8. Invoice Review and Payment")
add_bullets([
    "All invoices are subject to review and adjustment by the Department.",
    "The Company reserves the right to reduce or reject charges that do not comply with these guidelines.",
    "Payment terms are net 60 days from receipt of a compliant invoice.",
    "Disputes regarding invoice adjustments will be resolved through discussion between the General Counsel and the responsible partner.",
])

add_sub_header("9. Conflicts and Confidentiality")
add_bullets([
    "Outside counsel must disclose any actual or potential conflicts of interest prior to engagement.",
    "All Company information must be treated as confidential and may not be disclosed without written authorization.",
    "Outside counsel may not issue press releases, publish articles, or make public statements regarding Company matters without prior written consent.",
])

# ============================================================
# SAVE
# ============================================================
output_path = os.path.expanduser("~/Downloads/Legal Department Operations Manual - Complete.docx")
doc.save(output_path)
print(f"Document saved: {output_path}")
